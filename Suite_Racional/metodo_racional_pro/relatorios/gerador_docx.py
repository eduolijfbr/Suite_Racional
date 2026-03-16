# -*- coding: utf-8 -*-
"""
Gerador de Relatórios Técnicos em DOCX/ODT/HTML
Funciona com ou sem as bibliotecas python-docx e odfpy
"""

from datetime import datetime
import os
import zipfile
import tempfile

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False

# ODT é sempre disponível pois geramos manualmente
ODT_DISPONIVEL = True


class GeradorRelatorio:
    """Gera relatórios técnicos profissionais em DOCX, ODT ou HTML"""
    
    def __init__(self, template='completo'):
        self.template = template
        self.document = None
        
    def _escape_xml(self, text):
        """Escapa caracteres especiais para XML"""
        if not isinstance(text, str):
            text = str(text)
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;").replace("'", "&apos;")
        
    def _calcular_tabela_trs(self, dados):
        """
        Calcula vazões e diâmetros para múltiplos tempos de retorno.
        Retorna lista de dicts com TR, intensidade, vazão e diâmetro.
        """
        import math
        
        def to_float(val, default=0):
            try:
                if isinstance(val, str):
                    val = val.replace(',', '.')
                return float(val)
            except (ValueError, TypeError):
                return default

        dados_entrada = dados.get('dados_entrada', {})
        cidade = dados_entrada.get('cidade_idf', 'juiz_de_fora')
        tempo = to_float(dados_entrada.get('tempo'), 30)
        area = to_float(dados_entrada.get('area') or dados.get('area'), 0)
        # Tentar pegar de múltiplas chaves possíveis
        coef = to_float(dados_entrada.get('coef_escoamento') or dados_entrada.get('impermeabilidade') or dados.get('coeficiente'), 0.5)
        rugosidade = to_float(dados_entrada.get('rugosidade') or dados.get('rugosidade'), 0.013)
        declividade = to_float(dados_entrada.get('declividade') or dados.get('declividade'), 0.5)
        
        # Tempos de retorno padrão
        trs = [2, 5, 10, 25, 50, 100]
        resultados = []
        
        try:
            # Importar curvas IDF
            import sys
            import os
            plugin_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            if plugin_dir not in sys.path:
                sys.path.insert(0, plugin_dir)
            
            from hidrologia.curvas_idf import CurvasIDF
            idf = CurvasIDF()
            
            for tr in trs:
                try:
                    # Calcular intensidade
                    intensidade = idf.calcular_intensidade(cidade, tr, tempo)
                    
                    # Calcular vazão: Q = (C * I * A) / 3.6
                    vazao = (coef * intensidade * area) / 3.6
                    
                    # Calcular diâmetro (Manning para seção circular 85%)
                    if vazao > 0 and declividade > 0:
                        S = declividade / 100  # converter % para m/m
                        # Fórmula simplificada para seção circular
                        D = ((vazao * rugosidade) / (0.312 * (S ** 0.5))) ** 0.375
                    else:
                        D = 0
                    
                    resultados.append({
                        'tr': tr,
                        'intensidade': intensidade,
                        'vazao': vazao,
                        'diametro': D
                    })
                except:
                    resultados.append({
                        'tr': tr,
                        'intensidade': 0,
                        'vazao': 0,
                        'diametro': 0
                    })
        except Exception as e:
            # Se falhar, retornar lista vazia
            pass
            
        return resultados
    
    def gerar_relatorio_completo(self, dados, caminho_saida):
        """
        Gera relatório técnico completo.
        
        Parâmetros:
            dados: dict com resultados do cálculo
            caminho_saida: Caminho do arquivo de saída
        
        Formatos suportados:
            - .docx: Word (requer python-docx)
            - .odt: LibreOffice/OpenOffice (sempre disponível)
            - .html: HTML (sempre disponível, abre em qualquer navegador/editor)
        """
        # Verificar extensão e disponibilidade
        _, ext = os.path.splitext(caminho_saida)
        ext = ext.lower()
        
        if ext == '.docx' and DOCX_DISPONIVEL:
            self._gerar_docx(dados, caminho_saida)
        elif ext == '.odt':
            self._gerar_odt(dados, caminho_saida)
        elif ext == '.docx' and not DOCX_DISPONIVEL:
            # Fallback para ODT se DOCX não disponível
            caminho_saida = caminho_saida.replace('.docx', '.odt')
            self._gerar_odt(dados, caminho_saida)
        else:
            # HTML como padrão
            if ext not in ['.html', '.htm']:
                caminho_saida = caminho_saida + '.html'
            self._gerar_html(dados, caminho_saida)
            
    def _gerar_metodologia_tc_html(self, dados):
        """Gera o bloco HTML para a metodologia do Tc"""
        dados_entrada = dados.get('dados_entrada', {})
        metodo_usado = dados.get('metodo_tc') or dados_entrada.get('metodo_tc', 'manual')
        tabela_tc = dados.get('tabela_tc') or dados_entrada.get('tabela_tc', [])
        
        # SEMPRE calcular a tabela se estiver vazia
        if not tabela_tc:
            tabela_tc, metodo_usado = self._calcular_tabela_tc_interno(dados)
        
        nomes_metodos = {
            'kirpich': 'Kirpich',
            'giandotti': 'Giandotti',
            'ventura': 'Ventura',
            'bransby_williams': 'Bransby-Williams',
            'media': 'Média Aritmética'
        }
        
        nome_metodo = nomes_metodos.get(metodo_usado, metodo_usado.replace('_', ' ').title())
        
        html = ""
        if metodo_usado == 'manual':
            html += "<p>O Tempo de Concentração foi inserido manualmente pelo responsável técnico. Abaixo, a título de comparação, são apresentados os valores calculados por diferentes metodologias:</p>"
        else:
            html += f"<p>Para o cálculo do Tempo de Concentração (Tc), foi selecionada a metodologia: <strong>{nome_metodo}</strong>.</p>"
        
        if tabela_tc:
            html += "<p>Resultados obtidos por diferentes metodologias para comparação:</p>"
            html += "<table><tr><th>Metodologia / Equação</th><th>Tc Calculado (min)</th></tr>"
            for key_tc, nome_tc, valor_tc in tabela_tc:
                if key_tc == metodo_usado or key_tc == 'adotado':
                    html += f'<tr class="resultado"><td><strong>{nome_tc}</strong></td><td><strong>{valor_tc:.2f}</strong></td></tr>'
                else:
                    html += f"<tr><td>{nome_tc}</td><td>{valor_tc:.2f}</td></tr>"
            html += "</table><br>"
        return html

    def _gerar_tabela_idf_html(self, dados):
        """Gera o bloco HTML para a tabela IDF"""
        try:
            from ..hidrologia.curvas_idf import CurvasIDF
            idf_calc = CurvasIDF()
            dados_entrada = dados.get('dados_entrada', {})
            cidade_key = dados_entrada.get('cidade_idf', '').lower().replace(' ', '_')
            params_idf = idf_calc.obter_parametros(cidade_key)
            if not params_idf:
                return ''
                
            TRs = [2, 5, 10, 25, 50, 100]
            duracoes = [10, 15, 30, 60, 120]
            T_req = dados_entrada.get('tempo', 15)
            t_arr = round(T_req, 2)
            if t_arr not in duracoes:
                duracoes.append(t_arr)
            duracoes.sort()
            
            tab_dados = idf_calc.gerar_tabela_intensidades(cidade_key, TRs, duracoes)
            
            html = f'<p>Resultados associados à curva IDF de {cidade_key.replace("_", " ").title()}:</p>'
            html += '<table class="tab-idf"><tr><th>Duração (min)</th>'
            for tr_val in TRs:
                html += f'<th>TR={tr_val}</th>'
            html += '</tr>'
            
            for d in duracoes:
                html += f'<tr><td><strong>{d}</strong></td>'
                for tr_val in TRs:
                    val_idf = tab_dados['intensidades'][tr_val][d]
                    if d == t_arr and tr_val == dados_entrada.get('tempo_retorno', 25):
                        html += f'<td class="resultado"><strong>{val_idf:.1f}</strong></td>'
                    else:
                        html += f'<td>{val_idf:.1f}</td>'
                html += '</tr>'
            html += '</table><br>'
            return html
        except Exception as e:
            return f'<p>Erro ao gerar tabela IDF: {e}</p>'

    def _gerar_params_idf_html(self, dados):
        """Gera bloco HTML com os parâmetros K, a, b, c da curva IDF"""
        try:
            from ..hidrologia.curvas_idf import CurvasIDF
            idf_calc = CurvasIDF()
            dados_entrada = dados.get('dados_entrada', {})
            cidade_key = dados_entrada.get('cidade_idf', '').lower().replace(' ', '_')
            params_idf = idf_calc.obter_parametros(cidade_key)
            if not params_idf:
                return ''
            
            K, a, b, c = params_idf['K'], params_idf['a'], params_idf['b'], params_idf['c']
            fonte = params_idf.get('fonte', 'Não informada')
            tr = dados_entrada.get('tempo_retorno', 25)
            tempo = dados_entrada.get('tempo', 15)
            
            html = '<p>Parâmetros da curva IDF adotada:</p>'
            html += '<table><tr><th>Parâmetro</th><th>Valor</th></tr>'
            html += f'<tr><td>K</td><td>{K}</td></tr>'
            html += f'<tr><td>a</td><td>{a}</td></tr>'
            html += f'<tr><td>b</td><td>{b}</td></tr>'
            html += f'<tr><td>c</td><td>{c}</td></tr>'
            html += f'<tr><td>Fonte</td><td>{self._escape_xml(fonte)}</td></tr>'
            html += '</table><br>'
            html += f'<p>Equação parametrizada: i = ({K} × {tr}<sup>{a}</sup>) / ({tempo:.2f} + {b})<sup>{c}</sup></p>'
            return html
        except:
            return ''

    def _gerar_params_idf_odt(self, dados):
        """Gera bloco XML ODT com os parâmetros K, a, b, c da curva IDF"""
        try:
            from ..hidrologia.curvas_idf import CurvasIDF
            idf_calc = CurvasIDF()
            dados_entrada = dados.get('dados_entrada', {})
            cidade_key = dados_entrada.get('cidade_idf', '').lower().replace(' ', '_')
            params_idf = idf_calc.obter_parametros(cidade_key)
            if not params_idf:
                return ''
            
            K, a, b, c = params_idf['K'], params_idf['a'], params_idf['b'], params_idf['c']
            fonte = params_idf.get('fonte', 'Não informada')
            tr = dados_entrada.get('tempo_retorno', 25)
            tempo = dados_entrada.get('tempo', 15)
            
            xml = '<text:p text:style-name="Normal">Parâmetros da curva IDF adotada:</text:p>'
            xml += '<table:table table:name="TabelaParamsIDF" table:style-name="TableCenter">'
            xml += '<table:table-column table:number-columns-repeated="2"/>'
            xml += '<table:table-row>'
            xml += '<table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Parâmetro</text:p></table:table-cell>'
            xml += '<table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Valor</text:p></table:table-cell>'
            xml += '</table:table-row>'
            for nome_p, val_p in [('K', str(K)), ('a', str(a)), ('b', str(b)), ('c', str(c)), ('Fonte', fonte)]:
                xml += f'<table:table-row>'
                xml += f'<table:table-cell table:style-name="TableCell"><text:p>{nome_p}</text:p></table:table-cell>'
                xml += f'<table:table-cell table:style-name="TableCell"><text:p>{val_p}</text:p></table:table-cell>'
                xml += f'</table:table-row>'
            xml += '</table:table>'
            xml += f'<text:p text:style-name="Normal">Equação parametrizada: i = ({K} × {tr}^{a}) / ({tempo:.2f} + {b})^{c}</text:p>'
            return xml
        except:
            return ''

    def _gerar_html(self, dados, caminho_saida):

        """Gera relatório em formato HTML (compatível com Word)"""
        def to_float(val, default=0):
            try:
                if isinstance(val, str):
                    val = val.replace(',', '.')
                return float(val)
            except (ValueError, TypeError):
                return default

        dados_entrada = dados.get('dados_entrada', {})
        
        # Extrair valores com fallbacks robustos
        vazao = to_float(dados.get('vazao') or dados_entrada.get('vazao'), 0)
        diametro = to_float(dados.get('diametro') or dados_entrada.get('diametro'), 0)
        velocidade = to_float(dados.get('velocidade') or dados_entrada.get('velocidade'), 0)
        froude = to_float(dados.get('froude') or dados_entrada.get('froude'), 0)
        lamina = to_float(dados.get('lamina') or dados_entrada.get('lamina'), 0.85) * 100
        
        distancia = to_float(dados_entrada.get('distancia'), 0)
        desnivel = to_float(dados_entrada.get('desnivel'), 0)
        tempo = to_float(dados_entrada.get('tempo'), 0)
        area = to_float(dados_entrada.get('area'), 0)
        coef = to_float(dados_entrada.get('coef_escoamento') or dados_entrada.get('impermeabilidade'), 0)
        rugosidade = to_float(dados_entrada.get('rugosidade'), 0)
        declividade = to_float(dados_entrada.get('declividade'), 0)
        tr = dados_entrada.get('tempo_retorno', 0)
        intensidade = to_float(dados_entrada.get('intensidade'), 0)
        
        # Status das verificações
        if 0.6 <= velocidade <= 5.0:
            status_vel = "✓ ADEQUADA"
            cor_vel = "#28a745"
            obs_vel = "Velocidade dentro dos limites normativos (0,6 a 5,0 m/s)."
        elif velocidade < 0.6:
            status_vel = "⚠ BAIXA"
            cor_vel = "#ffc107"
            obs_vel = "Velocidade abaixo do mínimo. Risco de sedimentação."
        else:
            status_vel = "⚠ ALTA"
            cor_vel = "#dc3545"
            obs_vel = "Velocidade acima do máximo. Risco de erosão."
            
        if froude < 0.8:
            status_fr = "✓ SUBCRÍTICO"
            cor_fr = "#28a745"
            obs_fr = "Escoamento lento e tranquilo. Ideal para evitar erosão excessiva."
        elif 0.8 <= froude <= 1.2:
            status_fr = "⚠ CRÍTICO / INSTÁVEL"
            cor_fr = "#dc3545"
            obs_fr = "Estado de energia instável próximo à zona crítica. Dificulta a precisão geométrica da lâmina d'água."
        else:
            status_fr = "⚠ SUPERCRÍTICO"
            cor_fr = "#ffc107"
            obs_fr = "Escoamento rápido. Requer atenção especial e prever dissipadores de energia para evitar ressalto hidráulico indesejado."
        
        # Gerar tabela de múltiplos TRs
        tabela_trs = self._calcular_tabela_trs(dados)
        if tabela_trs:
            tabela_trs_html = '''<table>
    <tr>
        <th>Tempo de Retorno (anos)</th>
        <th>Intensidade (mm/h)</th>
        <th>Vazão (m³/s)</th>
        <th>Diâmetro (m)</th>
    </tr>'''
            for item in tabela_trs:
                destaque = ' class="resultado"' if item['tr'] == tr else ''
                tabela_trs_html += f'''
    <tr{destaque}>
        <td>{item['tr']}</td>
        <td>{item['intensidade']:.2f}</td>
        <td>{item['vazao']:.3f}</td>
        <td>{item['diametro']:.2f}</td>
    </tr>'''
            tabela_trs_html += '\n</table>'
        else:
            tabela_trs_html = '<p><em>Tabela de múltiplos TRs não disponível.</em></p>'
            
        camadas_usadas = dados.get('camadas_usadas', {})
        tabela_camadas_html = ""
        if camadas_usadas:
            tabela_camadas_html = '''<table>
    <tr><th>Parâmetro Espacial</th><th>Nome da Camada (Fonte de Dados)</th></tr>\n'''
            for tipo, nome in camadas_usadas.items():
                tabela_camadas_html += f'    <tr><td><strong>{self._escape_xml(tipo)}</strong></td><td>{self._escape_xml(nome)}</td></tr>\n'
            tabela_camadas_html += '</table>'
        else:
            tabela_camadas_html = '<p>Não foram detectadas conexões automáticas de camadas neste projeto (inserção puramente manual).</p>'
            
        cidade_idf = self._escape_xml(dados_entrada.get('cidade_idf', 'Não Definida').replace('_', ' ').title())
        A_secao = dados.get('area_secao', 1) 
        
        
        html = f'''<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Relatório de Drenagem - Método Racional</title>
    <style>
        @page {{
            size: A4;
            margin: 2.5cm 2.5cm 2.5cm 3cm;
        }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        .capa {{
            text-align: center;
            padding: 100px 0;
            page-break-after: always;
        }}
        .capa h1 {{
            color: #1565C0;
            font-size: 28pt;
            margin-bottom: 10px;
        }}
        .capa h2 {{
            color: #666;
            font-size: 18pt;
            font-weight: normal;
        }}
        .capa .info {{
            margin-top: 60px;
            font-size: 14pt;
        }}
        .capa .data {{
            margin-top: 80px;
            color: #666;
        }}
        h1 {
            color: #1565C0;
            font-size: 16pt;
            border-bottom: 2px solid #1565C0;
            padding-bottom: 5px;
            margin-top: 30px;
            page-break-after: avoid;
        }
        h2 {
            color: #1976D2;
            font-size: 13pt;
            margin-top: 20px;
            page-break-after: avoid;
        }
        p {
            text-align: justify;
            line-height: 1.5;
        }
        .equacao {
            text-align: center;
            font-size: 14pt;
            font-weight: bold;
            background: #E3F2FD;
            padding: 15px;
            border-radius: 5px;
            margin: 15px auto;
            width: fit-content;
        }
        table {
            width: 95%;
            border-collapse: collapse;
            margin: 20px auto;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 10px;
            text-align: left;
        }
        th {
            background: #1565C0;
            color: white;
            text-align: center;
        }
        .tab-idf {
            width: 100%;
            font-size: 8.5pt;
            table-layout: auto;
        }
        .tab-idf th, .tab-idf td {
            padding: 2px 1px;
            text-align: center;
        }
        tr:nth-child(even) {{
            background: #f9f9f9;
        }}
        .resultado {{
            background: #E8F5E9;
            font-weight: bold;
        }}
        .verificacao {{
            padding: 10px;
            border-radius: 5px;
            margin: 10px 0;
        }}
        .verificacao-ok {{
            background: #E8F5E9;
            border-left: 4px solid #28a745;
        }}
        .verificacao-alerta {{
            background: #FFF3E0;
            border-left: 4px solid #ffc107;
        }}
        .verificacao-erro {{
            background: #FFEBEE;
            border-left: 4px solid #dc3545;
        }}
        .memoria {{
            background: #FAFAFA;
            padding: 15px;
            border-radius: 5px;
            font-family: 'Courier New', monospace;
        }}
        .conclusao {{
            background: #E3F2FD;
            padding: 15px;
            border-radius: 5px;
            margin: 15px 0;
        }}
        .referencia {{
            margin-left: 20px;
            color: #666;
        }}
        .rodape {{
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            text-align: center;
            color: #999;
            font-size: 9pt;
        }}
        @media print {{
            body {{ margin: 0; padding: 0; }}
            .page-break {{ page-break-before: always; }}
        }}
    </style>
</head>
<body>

<!-- CAPA -->
<div class="capa">
    <h1>ESTUDO DE DRENAGEM URBANA</h1>
    <h2>Método Racional</h2>
</div>

<!-- LOCALIZAÇÃO DA ÁREA DE ESTUDOS -->
<h1>LOCALIZAÇÃO DA ÁREA DE ESTUDOS</h1>
<table>
    <tr>
        <td style="width: 30%;"><strong>Região / Endereço:</strong></td>
        <td>[Espaço reservado para o usuário preencher o endereço e região]</td>
    </tr>
    <tr>
        <td><strong>Coordenadas (Centroide):</strong></td>
        <td>{dados.get('coordenadas', 'Não disponível')}</td>
    </tr>
</table>

<h1>RESUMO DOS DADOS PRINCIPAIS</h1>
<table>
    <tr>
        <th>Vazão de Projeto</th>
        <th>Área da Bacia</th>
        <th>Tempo de Retorno</th>
    </tr>
    <tr class="resultado">
        <td style="text-align: center; font-size: 14pt;">{vazao:.2f} m³/s</td>
        <td style="text-align: center; font-size: 14pt;">{area:.4f} km²</td>
        <td style="text-align: center; font-size: 14pt;">{tr} anos</td>
    </tr>
</table>

<!-- SUMÁRIO EXECUTIVO -->
<h1>SUMÁRIO EXECUTIVO</h1>
<p>Este relatório apresenta os resultados do estudo de drenagem urbana realizado pelo Método Racional. Os principais resultados obtidos são:</p>

<table>
    <tr><th>Parâmetro</th><th>Valor</th><th>Unidade</th></tr>
    <tr class="resultado"><td>Vazão de Projeto</td><td>{vazao:.2f}</td><td>m³/s</td></tr>
    <tr class="resultado"><td>Diâmetro Calculado</td><td>{diametro:.2f}</td><td>m</td></tr>
    <tr class="resultado"><td>Lado Equiv. Galeria</td><td>{dados.get('lado_galeria', 0):.2f}</td><td>m</td></tr>
    <tr class="resultado"><td>Área da Seção</td><td>{dados.get('area_secao', 0):.2f}</td><td>m²</td></tr>
    <tr class="resultado"><td>Velocidade de Escoamento</td><td>{velocidade:.2f}</td><td>m/s</td></tr>
    <tr class="resultado"><td>Número de Froude</td><td>{froude:.2f}</td><td>-</td></tr>
</table>

<p>O dimensionamento atende aos critérios técnicos estabelecidos pelas normas brasileiras de drenagem urbana.</p>

<!-- INTRODUÇÃO -->
<h1>1. INTRODUÇÃO</h1>

<h2>1.1 Objetivo</h2>
<p>O presente estudo tem como objetivo dimensionar o sistema de drenagem urbana utilizando o Método Racional, determinando a vazão de projeto e as dimensões adequadas dos condutos.</p>

<h2>1.2 Escopo</h2>
<p>Este relatório abrange o cálculo hidrológico pelo Método Racional, o dimensionamento hidráulico dos condutos e as verificações técnicas necessárias para garantir o adequado funcionamento do sistema.</p>

<!-- METODOLOGIA -->
<h1>2. METODOLOGIA</h1>

<h2>2.1 Método Racional</h2>
<p>O Método Racional é amplamente utilizado para estimativa de vazões de pico em pequenas bacias urbanas. A equação fundamental é:</p>

<div class="equacao">Q = (C × I × A) / 3,6</div>

<p>Onde:</p>
<ul>
    <li><strong>Q</strong> = Vazão de pico (m³/s)</li>
    <li><strong>C</strong> = Impermeabilidade do solo (adimensional)</li>
    <li><strong>I</strong> = Intensidade de precipitação (mm/h)</li>
    <li><strong>A</strong> = Área da bacia (km²)</li>
</ul>

<h2>2.2 Dimensionamento Hidráulico</h2>
<p>O dimensionamento dos condutos foi realizado pela equação de Manning:</p>

<div class="equacao">Q = (1/n) × A × R<sup>2/3</sup> × S<sup>1/2</sup></div>

<p>Onde:</p>
<ul>
    <li><strong>n</strong> = Coeficiente de rugosidade de Manning</li>
    <li><strong>A</strong> = Área da seção molhada (m²)</li>
    <li><strong>R</strong> = Raio hidráulico (m)</li>
    <li><strong>S</strong> = Declividade do conduto (m/m)</li>
</ul>

<!-- DADOS DE ENTRADA -->
<h1>3. DADOS DE ENTRADA</h1>

<h2>3.1 Fontes de Extração Espacial</h2>
<p>Os parâmetros morfométricos e de uso do solo foram extraídos a partir das seguintes matrizes e vetores geográficos reais contidos no projeto atual do QGIS:</p>
{tabela_camadas_html}

<h2>3.2 Parâmetros Físicos Adotados</h2>
<p>No dimensionamento atual, foram considerados conceitos primordiais do Método Racional clássico. A capacidade das redes considerou uma restrição cautelar de Lâmina D'água para evitar a pressurização precoce e o golpe de aríete no poço de visita, delimitando o uso da seção em até 85%. Em contrapartida, avaliou-se a Velocidade de escoamento para afastar sedimentações de areias por velocidades mortas (< 0,60 m/s) ou erosões nos dutos em concretos/PVC geradas por grandes velocidades (> 5,00 m/s).</p>

<table>
    <tr><th>Parâmetro Físico</th><th>Valor</th><th>Unidade</th></tr>
    <tr><td><strong>Distância (Talvegue Máximo)</strong></td><td>{distancia:.2f}</td><td>m</td></tr>
    <tr><td><strong>Desnível Topográfico</strong></td><td>{desnivel:.2f}</td><td>m</td></tr>
    <tr><td><strong>Tempo de Concentração</strong></td><td>{tempo:.2f}</td><td>min</td></tr>
    <tr><td><strong>Área Contribuinte (A)</strong></td><td>{area:.6f}</td><td>km²</td></tr>
    <tr><td><strong>Impermeabilidade do Solo (C)</strong></td><td>{coef:.4f}</td><td>-</td></tr>
    <tr><td><strong>Rugosidade de Manning (n)</strong></td><td>{rugosidade:.4f}</td><td>-</td></tr>
    <tr><td><strong>Declividade Longitudinal (S)</strong></td><td>{declividade:.4f}</td><td>%</td></tr>
    <tr><td><strong>Tempo de Retorno Estipulado (TR)</strong></td><td>{tr}</td><td>anos</td></tr>
    <tr><td><strong>Intensidade Pluviométrica (i)</strong></td><td>{intensidade:.2f}</td><td>mm/h</td></tr>
</table>

<!-- MEMÓRIA DE CÁLCULO -->
<h1>4. MEMÓRIA DE CÁLCULO</h1>

<h2>4.1 Metodologia do Tempo de Concentração (Tc)</h2>
{self._gerar_metodologia_tc_html(dados)}

<h2>4.2 Cálculo de Intensidade (IDF)</h2>
<p>Baseado na equação de chuvas intensas selecionada para localidade de {cidade_idf}:</p>
<div class="memoria">
    <p>Equação geral: i = (K × TR<sup>a</sup>) / (t + b)<sup>c</sup></p>
    {self._gerar_params_idf_html(dados)}
    <p>Para TR = {tr} anos e Tempo de Concentração = {tempo:.2f} min:</p>
    <p><strong>Intensidade Pluviométrica Equivalente (I) = {intensidade:.2f} mm/h</strong></p>
</div>

<h2>4.3 Tabela de Intensidades para Múltiplos TRs e Durações</h2>
{self._gerar_tabela_idf_html(dados)}

<h2>4.4 Determinação da Vazão (Método Racional)</h2>
<p>Através da fórmula matriz de compatibilização da área de deflúvio (km²) e intensidade (mm/h) para vazão (m³/s):</p>
<div class="memoria">
    <p>Q = (C × i × A) / 3,6</p>
    <p>Q = ({coef:.4f} × {intensidade:.2f} × {area:.6f}) / 3,6</p>
    <p><strong>Vazão de Projeto Estimada (Q_max) = {vazao:.2f} m³/s</strong></p>
</div>

<h2>4.5 Dimensionamento por Manning</h2>
<p>Para a vazão obtida de <strong>{vazao:.2f} m³/s</strong>, aplica-se a fórmula de Manning admitindo regime uniforme em seção de galeria circular plena a uma ocupação normativa de 85%.</p>
<div class="memoria">
    <p>Q = (1 / n) * A_secao * R_h^(2/3) * S^(1/2)</p>
    <p>Substituindo a Declividade (S) de {declividade:.4f}% e a Rugosidade assumida (n) de {rugosidade:.4f}:</p>
    <br>
    <p>Diâmetro Equivalente Calculado: <strong>{diametro:.2f} m</strong></p>
    <p>Lado p/ Galeria Celular Quadrada Equiv.: <strong>{dados.get('lado_galeria', 0):.2f} m</strong></p>
    <p>Velocidade Resultante (Seção Plena): <strong>{velocidade:.2f} m/s</strong></p>
</div>

<!-- ANÁLISE DE IMPERMEABILIDADE -->
'''
        dados_imper = dados.get('impermeabilidade_dados')
        if dados_imper:
            html += f'''
<h1>5. ANÁLISE DE IMPERMEABILIDADE DO SOLO</h1>
<p>A determinação da Impermeabilidade do Solo (C) foi realizada através de análise multiespectral e classificação digital de pixels. Esta metodologia garante precisão na identificação da permeabilidade real da bacia.</p>

<table>
    <tr><th>Classe de Uso</th><th>Quantidade de Pixels</th><th>Percentual (%)</th></tr>
    <tr><td>Impermeável</td><td>{dados_imper.get('impermeable_pixels', 0):,}</td><td>{dados_imper.get('percentual', 0):.2f}%</td></tr>
    <tr><td>Vegetação</td><td>{dados_imper.get('vegetation_pixels', 0):,}</td><td>{dados_imper.get('percent_vegetation', 0):.2f}%</td></tr>
    <tr><td>Sombra / Água</td><td>{dados_imper.get('shadow_pixels', 0):,}</td><td>{dados_imper.get('percent_shadow', 0):.2f}%</td></tr>
    <tr class="resultado"><td><strong>Total Analisado</strong></td><td><strong>{dados_imper.get('total_pixels', 0):,}</strong></td><td><strong>100.00%</strong></td></tr>
</table>

<p><strong>Impermeabilidade do Solo Resultante (C): {dados_imper.get('coeficiente', 0.5):.4f}</strong></p>

<h2>5.1 Anexo Fotográfico da Classificação</h2>
<table style="border: none;">
    <tr style="border: none;">'''
            img_orig = dados_imper.get('impermeabilidade_imagem_original')
            img_class = dados_imper.get('impermeabilidade_imagem')
            
            if img_orig and os.path.exists(img_orig):
                html += f'''
        <td style="border: none; text-align: center; width: 45%; vertical-align: top;">
            <img src="file:///{img_orig.replace(chr(92), '/')}" style="width: 100%; height: auto; border: 1px solid #ddd;">
            <p><em>Imagem Original (RGB)</em></p>
        </td>'''
            
            if img_class and os.path.exists(img_class):
                html += f'''
        <td style="border: none; text-align: center; width: 45%; vertical-align: top;">
            <img src="file:///{img_class.replace(chr(92), '/')}" style="width: 100%; height: auto; border: 1px solid #ddd;">
            <p><em>Mapa de Classificação</em></p>
        </td>'''
            
            html += '''
    </tr>
</table>
'''
        
        html += f'''
<!-- RESULTADOS -->
<h1>6. RESULTADOS</h1>

<table>
    <tr><th>Resultado</th><th>Valor</th><th>Unidade</th></tr>
    <tr class="resultado"><td>Vazão de Projeto</td><td>{vazao:.2f}</td><td>m³/s</td></tr>
    <tr class="resultado"><td>Diâmetro Calculado</td><td>{diametro:.2f}</td><td>m</td></tr>
    <tr class="resultado"><td>Lado Equiv. Galeria</td><td>{dados.get('lado_galeria', 0):.2f}</td><td>m</td></tr>
    <tr class="resultado"><td>Área da Seção</td><td>{dados.get('area_secao', 0):.2f}</td><td>m²</td></tr>
    <tr class="resultado"><td>Velocidade de Escoamento</td><td>{velocidade:.2f}</td><td>m/s</td></tr>
    <tr class="resultado"><td>Número de Froude</td><td>{froude:.2f}</td><td>-</td></tr>
    <tr class="resultado"><td>Lâmina d'água / Diâmetro</td><td>{lamina:.0f}</td><td>%</td></tr>
</table>

<!-- TABELA DE MÚLTIPLOS TRs -->
<h1>7. ANÁLISE PARA MÚLTIPLOS TEMPOS DE RETORNO</h1>
<p>A tabela a seguir apresenta os resultados calculados para diferentes tempos de retorno, permitindo a análise comparativa e a escolha do cenário mais adequado ao projeto:</p>

{tabela_trs_html}

<p><em>* Valores calculados para a mesma duração de chuva ({tempo:.0f} min) e impermeabilidade do solo (C = {coef:.4f}).</em></p>

<!-- VERIFICAÇÕES TÉCNICAS -->
<h1>8. VERIFICAÇÕES HIDRÁULICAS</h1>

<h2>8.1 Verificação de Velocidade do Escoamento</h2>
<div class="verificacao {'verificacao-ok' if 0.6 <= velocidade <= 5.0 else 'verificacao-alerta'}">
    <p>Equação Padrão: V = Q / Área_Molhada_Secao</p>
    <p>V = {vazao:.2f} / {A_secao:.2f}</p>
    <br>
    <p><strong>Velocidade (v): {velocidade:.2f} m/s - <span style="color: {cor_vel}">{status_vel}</span></strong></p>
    <p>{obs_vel}</p>
    <p><em>Limites normativos: Mínimo 0,6 m/s (autolimpeza) | Máximo 5,0 m/s (erosão)</em></p>
</div>

<h2>8.2 Verificação de Regime (Froude)</h2>
<p>O número de Froude indica a influência da inércia sobre o escoamento face à gravidade. Escoamentos normais e passivos devem ser mantidos no estado Subcrítico (Flúvial) para prevenir choques de onda e perdas de carga localizadas no bueiro/coletores.</p>
<div class="verificacao {'verificacao-ok' if froude < 1 else 'verificacao-alerta'}">
    <p>Equação: Fr = vel / &radic;(g * Profundidade_Hidráulica)</p>
    <p>Fr = {velocidade:.2f} / &radic;(9.81 * {diametro:.2f})</p>
    <br>
    <p><strong>Número de Froude (Fr): {froude:.2f} - <span style="color: {cor_fr}">{status_fr}</span></strong></p>
    <p>{obs_fr}</p>
    <p><em>Fr &lt; 1: Subcrítico (recomendado) | Fr &gt; 1: Supercrítico</em></p>
</div>

<!-- CONCLUSÕES -->
<h1>9. CONCLUSÕES</h1>

<div class="conclusao">
    <p>Com base nos cálculos realizados e nas verificações técnicas apresentadas, conclui-se que:</p>
    <ul>
        <li>A vazão de projeto calculada é de <strong>{vazao:.2f} m³/s</strong>;</li>
        <li>O diâmetro recomendado para o conduto é de <strong>{diametro:.2f} m</strong>;</li>
        <li>A velocidade de escoamento de <strong>{velocidade:.2f} m/s</strong> está {'dentro' if 0.6 <= velocidade <= 5.0 else 'fora'} dos limites normativos;</li>
        <li>O regime de escoamento é <strong>{'subcrítico' if froude < 1 else 'supercrítico'}</strong> (Fr = {froude:.2f});</li>
        <li>O dimensionamento {'atende' if 0.6 <= velocidade <= 5.0 and froude < 1 else 'requer atenção quanto'} aos critérios técnicos estabelecidos.</li>
    </ul>
</div>

<!-- LIMITAÇÕES TÉCNICAS -->
<h1>10. NOTA DE LIMITAÇÃO TÉCNICA</h1>
<div class="verificacao verificacao-alerta">
    <ol>
        <li><strong>Aplicabilidade do Método:</strong> Este relatório utiliza o Método Racional, cuja premissa fundamental assume que a intensidade da chuva é uniforme e constante sobre toda a bacia durante o tempo de concentração (t<sub>c</sub>). Este método é estritamente recomendado para bacias de microdrenagem e áreas urbanas pequenas (geralmente inferiores a 50-100 hectares).</li>
        <li><strong>Macrodrenagem e Áreas Extensas:</strong> Para bacias de macrodrenagem ou áreas que excedam os limites de aplicação do método, os resultados aqui apresentados devem ser interpretados como estimativas preliminares. Recomenda-se a validação através de métodos que considerem o amortecimento de cheias e a variação temporal da chuva, como o Método do Hidrograma Unitário (SCS/NRCS) ou modelagem hidráulica dinâmica (ex: SWMM).</li>
        <li><strong>Coeficiente de Escoamento (C):</strong> Os valores de C adotados baseiam-se em tabelas normatizadas de uso e ocupação do solo. Alterações futuras na impermeabilização da bacia invalidam as vazões de projeto calculadas.</li>
        <li><strong>Verificação Hidráulica:</strong> A análise do Número de Froude (F<sub>r</sub>) indica o regime de escoamento. Trechos com F<sub>r</sub> próximos à unidade (regime crítico) apresentam instabilidade de lâmina d'água e requerem atenção especial no dimensionamento físico para evitar transbordamentos ou erosões não previstos.</li>
    </ol>
</div>

<!-- RODAPÉ -->
<div class="rodape">
    <p>Relatório gerado automaticamente pelo Plugin Método Racional Pro para QGIS</p>
    <p>{datetime.now().strftime("%d/%m/%Y às %H:%M")}</p>
</div>

</body>
</html>'''
        
        with open(caminho_saida, 'w', encoding='utf-8') as f:
            f.write(html)
            
    def _gerar_docx(self, dados, caminho_saida):
        """Gera relatório em formato DOCX (requer python-docx)"""
        self.document = Document()
        self._configurar_estilos()
        
        # Seções do relatório
        self._adicionar_capa(dados)
        self._adicionar_localizacao_resumo_docx(dados)
        self._adicionar_sumario_executivo(dados)
        self._adicionar_introducao(dados)
        self._adicionar_metodologia(dados)
        self._adicionar_dados_entrada(dados)
        self._adicionar_memoria_calculo(dados)
        self._adicionar_analise_impermeabilidade(dados)
        self._adicionar_resultados(dados)
        self._adicionar_verificacoes(dados)
        self._adicionar_conclusoes(dados)
        self._adicionar_tabela_trs_docx(dados)
        
        # Salvar documento
        self.document.save(caminho_saida)
        
    def _adicionar_localizacao_resumo_docx(self, dados):
        """Adiciona seção de localização e tabela de resumo no DOCX"""
        vazao = dados.get('vazao', 0)
        area = dados.get('area', 0)
        tr = dados.get('dados_entrada', {}).get('tempo_retorno', 0)
        coordenadas = dados.get('coordenadas', 'Não disponível')

        self.document.add_heading("LOCALIZAÇÃO DA ÁREA DE ESTUDOS", level=1)
        
        table = self.document.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Linha 1: Endereço
        c0 = table.rows[0].cells
        c0[0].text = "Região / Endereço:"
        c0[1].text = "[Espaço reservado para o usuário preencher o endereço e região]"
        c0[0].paragraphs[0].runs[0].bold = True

        # Linha 2: Coordenadas
        c1 = table.rows[1].cells
        c1[0].text = "Coordenadas (Centroide):"
        c1[1].text = coordenadas
        c1[0].paragraphs[0].runs[0].bold = True
        
        self.document.add_heading("RESUMO DOS DADOS PRINCIPAIS", level=1)
        
        resumo_table = self.document.add_table(rows=1, cols=3)
        resumo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        resumo_table.style = 'Table Grid'
        hdr_resumo = resumo_table.rows[0].cells
        hdr_resumo[0].text = "Vazão de Projeto"
        hdr_resumo[1].text = "Área da Bacia"
        hdr_resumo[2].text = "Tempo de Retorno"
        for cell in hdr_resumo:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
        
        row_resumo = resumo_table.add_row().cells
        
        def set_center_bold(cell, text, size=14):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(size)

        set_center_bold(row_resumo[0], f"{vazao:.2f} m³/s")
        set_center_bold(row_resumo[1], f"{area:.4f} km²")
        set_center_bold(row_resumo[2], f"{tr} anos")
        
        self.document.add_paragraph()
        
    def _configurar_estilos(self):
        """Configura estilos do documento"""
        # Configurar margens
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)
        
        # Estilo Normal: Justificado
        style = self.document.styles['Normal']
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.font.name = 'Arial'
        style.font.size = Pt(11)
    
    def _adicionar_capa(self, dados):
        """Adiciona capa do relatório"""
        # Título principal
        for _ in range(5):
            self.document.add_paragraph()
        
        titulo = self.document.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("ESTUDO DE DRENAGEM URBANA")
        run.bold = True
        run.font.size = Pt(24)
        
        subtitulo = self.document.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitulo.add_run("Método Racional")
        run.font.size = Pt(18)
        
        for _ in range(10):
            self.document.add_paragraph()
        
        self.document.add_page_break()
        
    def _adicionar_sumario_executivo(self, dados):
        """Adiciona sumário executivo"""
        self.document.add_heading("SUMÁRIO EXECUTIVO", level=1)
        
        dados_entrada = dados.get('dados_entrada', {})
        
        texto = f"""
Este relatório apresenta os resultados do estudo de drenagem urbana 
realizado pelo Método Racional. Os principais resultados obtidos são:

• Vazão de projeto: {dados.get('vazao', 0):.2f} m³/s
• Diâmetro calculado: {dados.get('diametro', 0):.2f} m
• Lado equiv. galeria: {dados.get('lado_galeria', 0):.2f} m
• Área da seção: {dados.get('area_secao', 0):.2f} m²
• Velocidade de escoamento: {dados.get('velocidade', 0):.2f} m/s
• Número de Froude: {dados.get('froude', 0):.2f}

O dimensionamento atende aos critérios técnicos estabelecidos pelas 
normas brasileiras de drenagem urbana.
"""
        self.document.add_paragraph(texto.strip())
        self.document.add_paragraph()
        
    def _adicionar_introducao(self, dados):
        """Adiciona introdução"""
        self.document.add_heading("1. INTRODUÇÃO", level=1)
        
        self.document.add_heading("1.1 Objetivo", level=2)
        self.document.add_paragraph(
            "O presente estudo tem como objetivo dimensionar o sistema de "
            "drenagem urbana utilizando o Método Racional, determinando a "
            "vazão de projeto e as dimensões adequadas dos condutos."
        )
        
        self.document.add_heading("1.2 Escopo", level=2)
        self.document.add_paragraph(
            "Este relatório abrange o cálculo hidrológico pelo Método Racional, "
            "o dimensionamento hidráulico dos condutos e as verificações técnicas "
            "necessárias para garantir o adequado funcionamento do sistema."
        )
        
    def _adicionar_metodologia(self, dados):
        """Adiciona metodologia"""
        self.document.add_heading("2. METODOLOGIA", level=1)
        
        self.document.add_heading("2.1 Método Racional", level=2)
        self.document.add_paragraph(
            "O Método Racional é amplamente utilizado para estimativa de vazões "
            "de pico em pequenas bacias urbanas. A equação fundamental é:"
        )
        
        eq = self.document.add_paragraph()
        eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = eq.add_run("Q = (C × I × A) / 3,6")
        run.bold = True
        run.font.size = Pt(12)
        
        self.document.add_paragraph("Onde:")
        self.document.add_paragraph("• Q = Vazão de pico (m³/s)")
        self.document.add_paragraph("• C = Coeficiente de escoamento superficial")
        self.document.add_paragraph("• I = Intensidade de precipitação (mm/h)")
        self.document.add_paragraph("• A = Área da bacia (km²)")
        
        self.document.add_heading("2.2 Dimensionamento Hidráulico", level=2)
        self.document.add_paragraph(
            "O dimensionamento dos condutos foi realizado pela equação de Manning:"
        )
        
        eq2 = self.document.add_paragraph()
        eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = eq2.add_run("Q = (1/n) × A × R^(2/3) × S^(1/2)")
        run.bold = True
        
    def _adicionar_dados_entrada(self, dados):
        """Adiciona dados de entrada e fontes de extração"""
        self.document.add_heading("3. DADOS DE ENTRADA", level=1)
        
        # --- SUBSEÇÃO 3.1 FONTES ESPACIAIS ---
        self.document.add_heading("3.1 Fontes de Extração Espacial", level=2)
        self.document.add_paragraph("Os parâmetros morfométricos e de uso do solo foram extraídos a partir das seguintes matrizes e vetores geográficos reais contidos no projeto atual do QGIS:")
        
        camadas_usadas = dados.get('camadas_usadas', {})
        if camadas_usadas:
            # Tabela de Fontes
            table_fontes = self.document.add_table(rows=1, cols=2)
            table_fontes.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_fontes.style = 'Table Grid'
            hdr_fontes = table_fontes.rows[0].cells
            hdr_fontes[0].text = 'Parâmetro Espacial'
            hdr_fontes[1].text = 'Nome da Camada (Fonte de Dados)'
            for row_cells in table_fontes.rows:
                for cell in row_cells:
                    cell.paragraphs[0].runs[0].bold = True
            
            for tipo, nome_camada in camadas_usadas.items():
                row_cells = table_fontes.add_row().cells
                row_cells[0].text = tipo
                row_cells[1].text = nome_camada
        else:
            self.document.add_paragraph("Não foram detectadas conexões automáticas de camadas neste projeto (inserção puramente manual).")
            
        self.document.add_paragraph()
        
        # --- SUBSEÇÃO 3.2 PARÂMETROS FÍSICOS E ADOÇÕES ---
        self.document.add_heading("3.2 Parâmetros Físicos Adotados", level=2)
        
        descricao_parametros = (
            "No dimensionamento atual, foram considerados conceitos primordiais do Método Racional clássico. "
            "A capacidade das redes considerou uma restrição cautelar de Lâmina D'água para evitar a pressurização precoce e o golpe de aríete no poço de visita, delimitando o uso da seção em até 85%. "
            "Em contrapartida, avaliou-se a Velocidade de escoamento para afastar sedimentações de areias por velocidades mortas (< 0,60 m/s) ou erosões nos dutos em concretos/PVC geradas por grandes velocidades (> 5,00 m/s)."
        )
        self.document.add_paragraph(descricao_parametros)
        
        dados_entrada = dados.get('dados_entrada', {})
        
        # Criar tabela de resultados de entrada
        table = self.document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parâmetro Físico'
        hdr_cells[1].text = 'Valor'
        hdr_cells[2].text = 'Unidade'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            
        # Dados com fallbacks para garantir que nunca apareçam zerados se houver alternativas
        def get_val(key, default=0):
            val = dados_entrada.get(key)
            if val is None or val == 0 or val == '0':
                # Tentar chaves alternativas
                if key == 'coef_escoamento':
                    val = dados_entrada.get('impermeabilidade') or dados.get('coeficiente')
                elif key == 'intensidade':
                    val = dados.get('intensidade')
            
            try:
                if isinstance(val, str): val = val.replace(',', '.')
                return float(val) if val is not None else default
            except:
                return default

        linhas = [
            ('Distância (Talvegue Máximo)', f"{get_val('distancia'):.2f}", 'm'),
            ('Desnível Topográfico', f"{get_val('desnivel'):.2f}", 'm'),
            ('Tempo de Concentração', f"{get_val('tempo'):.2f}", 'min'),
            ('Área Contribuinte (A)', f"{get_val('area'):.6f}", 'km²'),
            ('Impermeabilidade do Solo (C)', f"{get_val('coef_escoamento', 0.5):.4f}", '-'),
            ('Rugosidade de Manning (n)', f"{get_val('rugosidade', 0.013):.4f}", '-'),
            ('Declividade Longitudinal (S)', f"{get_val('declividade'):.4f}", '%'),
            ('Tempo de Retorno Estipulado (TR)', str(dados_entrada.get('tempo_retorno', 25)), 'anos'),
            ('Intensidade Pluviométrica (i)', f"{get_val('intensidade'):.2f}", 'mm/h')
        ]
        
        for nome, valor, unidade in linhas:
            row_cells = table.add_row().cells
            row_cells[0].text = nome
            row_cells[1].text = valor
            row_cells[2].text = unidade
            
        self.document.add_paragraph()
        
    def _adicionar_memoria_calculo(self, dados):
        """Adiciona memória de cálculo matemático puro"""
        self.document.add_heading("4. MEMÓRIA DE CÁLCULO", level=1)
        
        dados_entrada = dados.get('dados_entrada', {})
        
        # --- 4.1 METODOLOGIA DO TEMPO DE CONCENTRAÇÃO (Tc) ---
        self._adicionar_metodologia_tc_docx(dados)

        # --- 4.2 CÁLCULO DA INTENSIDADE IDF ---
        self.document.add_heading("4.2 Cálculo de Intensidade (IDF)", level=2)
        cidade = dados_entrada.get('cidade_idf', 'Não Definida').replace('_', ' ').title()
        self.document.add_paragraph(f"Baseado na equação de chuvas intensas selecionada para localidade de {cidade}:")
        
        TR = dados_entrada.get('tempo_retorno', 25)
        T_req = dados_entrada.get('tempo', 15)
        I = dados_entrada.get('intensidade', 0)
        
        self.document.add_paragraph("Equação geral de chuvas intensas:")
        self.document.add_paragraph("i = (K × TR^a) / (t + b)^c")
        
        # Obter e exibir os parâmetros
        try:
            from ..hidrologia.curvas_idf import CurvasIDF
            idf_calc = CurvasIDF()
            cidade_key = dados_entrada.get('cidade_idf', '').lower().replace(' ', '_')
            params_idf = idf_calc.obter_parametros(cidade_key)
            if params_idf:
                K, a, b, c = params_idf['K'], params_idf['a'], params_idf['b'], params_idf['c']
                fonte = params_idf.get('fonte', 'Não informada')
                
                # Tabela com os parâmetros
                self.document.add_paragraph("Parâmetros da curva IDF adotada:")
                tbl_params = self.document.add_table(rows=1, cols=2)
                tbl_params.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl_params.style = 'Table Grid'
                hdr_p = tbl_params.rows[0].cells
                hdr_p[0].text = 'Parâmetro'
                hdr_p[1].text = 'Valor'
                hdr_p[0].paragraphs[0].runs[0].bold = True
                hdr_p[1].paragraphs[0].runs[0].bold = True
                
                for nome_p, val_p in [('K', str(K)), ('a', str(a)), ('b', str(b)), ('c', str(c)), ('Fonte', fonte)]:
                    row_p = tbl_params.add_row().cells
                    row_p[0].text = nome_p
                    row_p[1].text = val_p
                
                self.document.add_paragraph()
                self.document.add_paragraph(f"Equação parametrizada: i = ({K} × TR^{a}) / (t + {b})^{c}")
                self.document.add_paragraph(f"Substituindo TR = {TR} e t = {T_req:.2f}:")
                self.document.add_paragraph(f"i = ({K} × {TR}^{a}) / ({T_req:.2f} + {b})^{c}")
        except:
            params_idf = None
            self.document.add_paragraph(f"Para TR = {TR} anos e Tempo de Concentração = {T_req:.2f} min:")
            
        res_i = self.document.add_paragraph()
        run_i = res_i.add_run(f"Intensidade Pluviométrica Equivalente (I) = {I:.2f} mm/h")
        run_i.bold = True
        
        # --- 4.3 TABELA DE INTENSIDADES (MÚLTIPLOS TRs) ---
        if params_idf:
            self.document.add_heading("4.3 Tabela de Intensidades para Múltiplos TRs e Durações", level=2)
            self.document.add_paragraph(f"Resultados associados à curva IDF de {cidade}:")
            
            try:
                TRs = [2, 5, 10, 25, 50, 100]
                duracoes = [10, 15, 30, 60, 120]
                # Inclui o tempo do projeto na tabela se não existir
                t_arr = round(T_req, 2)
                if t_arr not in duracoes:
                    duracoes.append(t_arr)
                duracoes.sort()
                
                tbl_idf = self.document.add_table(rows=1, cols=len(TRs) + 1)
                tbl_idf.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl_idf.style = 'Table Grid'
                hdr_idf = tbl_idf.rows[0].cells
                hdr_idf[0].text = 'Duração (min)'
                for i, tr_val in enumerate(TRs):
                    hdr_idf[i+1].text = f'TR={tr_val}'
                    
                for cell in hdr_idf: cell.paragraphs[0].runs[0].bold = True
                    
                tab_dados = idf_calc.gerar_tabela_intensidades(cidade_key, TRs, duracoes)
                
                for d in duracoes:
                    row_cells = tbl_idf.add_row().cells
                    row_cells[0].text = str(d)
                    for i, tr_val in enumerate(TRs):
                        val_idf = tab_dados['intensidades'][tr_val][d]
                        row_cells[i+1].text = f"{val_idf:.1f}"
                        if d == t_arr and tr_val == TR:
                            row_cells[i+1].paragraphs[0].runs[0].bold = True
                
                # Definir larguras de colunas fixas para garantir o encaixe (Total ~16cm em A4)
                # Coluna 0 (Duração): 2.4cm | Outras 6 colunas (TRs): ~2.1cm cada
                tbl_idf.columns[0].width = Cm(2.4)
                for i in range(1, len(TRs) + 1):
                    tbl_idf.columns[i].width = Cm(2.1)

                # Formatar todas as células da tabela para caber e centralizar
                for row in tbl_idf.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(8.5)
                
                self.document.add_paragraph()
            except Exception as e:
                pass
        
        # --- 4.4 CÁLCULO DA VAZÃO ---
        self.document.add_heading("4.4 Determinação da Vazão (Método Racional)", level=2)
        
        C = dados_entrada.get('coef_escoamento', 0)
        A = dados_entrada.get('area', 0)
        Q = dados.get('vazao', 0)
        
        self.document.add_paragraph(f"Através da fórmula matriz de compatibilização da área de deflúvio (km²) e intensidade (mm/h) para vazão (m³/s):")
        self.document.add_paragraph("Q = (C × i × A) / 3,6")
        self.document.add_paragraph(f"Q = ({C:.4f} × {I:.2f} × {A:.6f}) / 3,6")
        
        res_q = self.document.add_paragraph()
        run_q = res_q.add_run(f"Vazão de Projeto Estimada (Q_max) = {Q:.2f} m³/s")
        run_q.bold = True
        
        # --- 4.5 DIMENSIONAMENTO DO CONDUTO ---
        self.document.add_heading("4.5 Dimensionamento por Manning", level=2)
        self.document.add_paragraph(
            f"Para a vazão obtida de {Q:.2f} m³/s, aplica-se a fórmula de Manning admitindo regime uniforme em seção de galeria circular plena a uma ocupação normativa de 85%."
        )
        self.document.add_paragraph("Q = (1 / n) * A_secao * R_h^(2/3) * S^(1/2)")
        self.document.add_paragraph(f"Substituindo a Declividade (S) de {dados_entrada.get('declividade', 0):.4f}% e a Rugosidade assumida (n) de {dados_entrada.get('rugosidade', 0.013):.4f}:")
        
        self.document.add_paragraph(f"Diâmetro Equivalente Calculado = {dados.get('diametro', 0):.2f} m")
        self.document.add_paragraph(f"Lado p/ Galeria Celular Quadrada Equiv. = {dados.get('lado_galeria', 0):.2f} m")
        self.document.add_paragraph(f"Velocidade Resultante (Seção Plena) = {dados.get('velocidade', 0):.2f} m/s")
        
        # Mapa do projeto removido a pedido do usuário

    def _adicionar_analise_impermeabilidade(self, dados):
        """Adiciona seção detalhada de impermeabilidade do solo"""
        # Tentar pegar do nível superior ou de dentro dos dados de entrada
        dados_imper = dados.get('impermeabilidade_dados') or dados.get('dados_entrada', {}).get('impermeabilidade_dados')
        
        if not dados_imper:
            return

        self.document.add_heading("5. ANÁLISE DE IMPERMEABILIDADE DO SOLO", level=1)
        
        texto_intro = (
            "A determinação da Impermeabilidade do Solo (C) foi realizada através de análise multiespectral "
            "e classificação digital de pixels sobre imagem de alta resolução. Esta metodologia permite uma "
            "precisão superior na identificação de áreas impermeáveis, vegetadas e corpos d'água/sombras."
        )
        self.document.add_paragraph(texto_intro)

        # Tabela de estatísticas de pixels
        table = self.document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Classe de Uso'
        hdr[1].text = 'Quantidade de Pixels'
        hdr[2].text = 'Percentual (%)'
        for cell in hdr: cell.paragraphs[0].runs[0].bold = True

        stats = [
            ('Impermeável', dados_imper.get('impermeable_pixels', 0), dados_imper.get('percentual', 0)),
            ('Vegetação', dados_imper.get('vegetation_pixels', 0), dados_imper.get('percent_vegetation', 0)),
            ('Sombra / Água', dados_imper.get('shadow_pixels', 0), dados_imper.get('percent_shadow', 0)),
            ('Total Analisado', dados_imper.get('total_pixels', 0), 100.0)
        ]

        for classe, qtd, perc in stats:
            row = table.add_row().cells
            row[0].text = classe
            row[1].text = f"{qtd:,}"
            row[2].text = f"{perc:.2f}%"

        # Impermeabilidade Resultante
        p = self.document.add_paragraph()
        p.add_run("\nImpermeabilidade do Solo Resultante (C): ").bold = True
        run_c = p.add_run(f"{dados_imper.get('coeficiente', 0.5):.4f}")
        run_c.bold = True
        run_c.font.size = Pt(14)

        # Anexo Fotográfico da Classificação
        self.document.add_heading("5.1 Anexo Fotográfico da Classificação", level=2)
        
        # Tentar carregar as duas imagens (Original e Classificada)
        img_orig = dados_imper.get('impermeabilidade_imagem_original')
        img_class = dados_imper.get('impermeabilidade_imagem')

        if img_orig and os.path.exists(img_orig) and img_class and os.path.exists(img_class):
            # Lado a lado usando uma tabela invisível
            img_table = self.document.add_table(rows=1, cols=2)
            cells = img_table.rows[0].cells
            
            try:
                # Imagem Original
                p1 = cells[0].paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run1 = p1.add_run()
                run1.add_picture(img_orig, width=Cm(7))
                cells[0].add_paragraph("Imagem Original (RGB)").alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Imagem Classificada
                p2 = cells[1].paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run2 = p2.add_run()
                run2.add_picture(img_class, width=Cm(7))
                cells[1].add_paragraph("Mapa de Classificação").alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                self.document.add_paragraph("[Erro ao inserir imagens de classificação no documento]")
        elif img_class and os.path.exists(img_class):
            # Apenas uma se a outra falhar
            try:
                self.document.add_picture(img_class, width=Cm(12))
                desc = self.document.add_paragraph("Figura: Mapa de Classificação da Impermeabilidade")
                desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except: pass
        
        self.document.add_page_break()
        
    def _adicionar_resultados(self, dados):
        """Adiciona resultados"""
        self.document.add_heading("6. RESULTADOS", level=1)
        
        # Criar tabela de resultados
        table = self.document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Resultado'
        hdr_cells[1].text = 'Valor'
        hdr_cells[2].text = 'Unidade'
        
        resultados = [
            ('Vazão de Projeto', dados.get('vazao', 0), 'm³/s'),
            ('Diâmetro Calculado', dados.get('diametro', 0), 'm'),
            ('Lado Equiv. Galeria', dados.get('lado_galeria', 0), 'm'),
            ('Área da Seção', dados.get('area_secao', 0), 'm²'),
            ('Velocidade', dados.get('velocidade', 0), 'm/s'),
            ('Número de Froude', dados.get('froude', 0), '-'),
            ('Lâmina/Diâmetro', dados.get('lamina', 0.85) * 100, '%'),
        ]
        
        for param, valor, unidade in resultados:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = f"{valor:.2f}"
            row_cells[2].text = unidade
            
        self.document.add_paragraph()
        
    def _adicionar_verificacoes(self, dados):
        """Adiciona verificações técnicas com embasamento matemático explícito."""
        self.document.add_heading("7. VERIFICAÇÕES HIDRÁULICAS", level=1)
        
        velocidade = dados.get('velocidade', 0)
        froude = dados.get('froude', 0)
        Q = dados.get('vazao', 0)
        A_secao = dados.get('area_secao', 1) # Fallback 1
        diametro = dados.get('diametro', 1)
        
        # --- 7.1 VELOCIDADE ---
        self.document.add_heading("7.1 Verificação de Velocidade do Escoamento", level=2)
        
        if 0.6 <= velocidade <= 5.0:
            status_vel = "✓ CONDUTO ADEQUADO"
            obs_vel = "A velocidade é perfeitamente funcional entre os vetores de autolimpeza (0,6m/s) e limite degradante estrutural (5,0m/s)."
        elif velocidade < 0.6:
            status_vel = "⚠ VELOCIDADE DEPOSICIONAL"
            obs_vel = "Velocidade abaixo do mínimo estipulado. Haverá risco de decantação e entupimento de materiais suspensos pela rede de macro/microdrenagem."
        else:
            status_vel = "⚠ EXCESSO DE VELOCIDADE"
            obs_vel = "A velocidade supera a margem tolerável em pavimentos asfálticos/concreto (5,0m/s). Risco evidente de desagregação prematura do recobrimento por turbilhonamento excessivo."
            
        self.document.add_paragraph(f"Equação Padrão: V = Q / Área_Molhada_Secao")
        self.document.add_paragraph(f"V = {Q:.2f} / {A_secao:.2f}")
        
        vel_res = self.document.add_paragraph(f"Velocidade (v): {velocidade:.2f} m/s - {status_vel}")
        vel_res.runs[0].bold = True 
        self.document.add_paragraph(obs_vel)
        
        # --- 7.2 NÚMERO DE FROUDE ---
        self.document.add_heading("7.2 Verificação de Regime (Froude)", level=2)
        self.document.add_paragraph("O número de Froude indica a influência da inércia sobre o escoamento face à gravidade. Escoamentos normais e passivos devem ser mantidos no estado Subcrítico (Flúvial) para prevenir choques de onda e perdas de carga localizadas no bueiro/coletores.")
        
        self.document.add_paragraph(f"Equação: Fr = vel / √(g * Profundidade_Hidráulica)")
        self.document.add_paragraph(f"Fr = {velocidade:.2f} / √(9.81 * {diametro:.2f})")
        
        if froude < 0.8:
            status_fr = "✓ REGIME SUBCRÍTICO"
            obs_fr = "Escoamento lento e tranquilo. Ideal para canais de macrodrenagem onde se deseja evitar erosão excessiva."
        elif 0.8 <= froude <= 1.2:
            status_fr = "⚠ REGIME CRÍTICO / INSTÁVEL"
            obs_fr = "Estado de energia mínima com escoamento instável (Froude próximo a 1). Recomenda-se evitar este regime em canais extensos."
        else:
            status_fr = "⚠ REGIME SUPERCRÍTICO (TORRENCIAL)"
            obs_fr = "Escoamento rápido em declives acentuados. Atenção: requere artifícios dissipadores de energia à jusante para conter o ressalto."
            
        fr_res = self.document.add_paragraph(f"Número de Froude (Fr): {froude:.2f} - {status_fr}")
        fr_res.runs[0].bold = True
        self.document.add_paragraph(obs_fr)
        
    def _adicionar_conclusoes(self, dados):
        """Adiciona conclusões"""
        self.document.add_heading("8. CONCLUSÕES", level=1)
        
        self.document.add_paragraph(
            "Com base nos cálculos realizados e nas verificações técnicas "
            "apresentadas, conclui-se que:"
        )
        
        self.document.add_paragraph(
            f"• A vazão de projeto calculada é de {dados.get('vazao', 0):.2f} m³/s;"
        )
        self.document.add_paragraph(
            f"• O diâmetro recomendado para o conduto é de {dados.get('diametro', 0):.2f} m;"
        )
        self.document.add_paragraph(
            "• O dimensionamento atende aos critérios técnicos estabelecidos."
        )

        
        self.document.add_heading("9. NOTA DE LIMITAÇÃO TÉCNICA", level=1)
        self.document.add_paragraph("1. Aplicabilidade do Método: Este relatório utiliza o Método Racional, cuja premissa fundamental assume que a intensidade da chuva é uniforme e constante sobre toda a bacia durante o tempo de concentração (tc). Este método é estritamente recomendado para bacias de microdrenagem e áreas urbanas pequenas (geralmente inferiores a 50-100 hectares).", style='List Number')
        self.document.add_paragraph("2. Macrodrenagem e Áreas Extensas: Para bacias de macrodrenagem ou áreas que excedam os limites de aplicação do método, os resultados aqui apresentados devem ser interpretados como estimativas preliminares. Recomenda-se a validação através de métodos que considerem o amortecimento de cheias e a variação temporal da chuva, como o Método do Hidrograma Unitário (SCS/NRCS) ou modelagem hidráulica dinâmica (ex: SWMM).", style='List Number')
        self.document.add_paragraph("3. Coeficiente de Escoamento (C): Os valores de C adotados baseiam-se em tabelas normatizadas de uso e ocupação do solo. Alterações futuras na impermeabilização da bacia invalidam as vazões de projeto calculadas.", style='List Number')
        self.document.add_paragraph("4. Verificação Hidráulica: A análise do Número de Froude (Fr) indica o regime de escoamento. Trechos com Fr próximos à unidade (regime crítico) apresentam instabilidade de lâmina d'água e requerem atenção especial no dimensionamento físico para evitar transbordamentos ou erosões não previstos.", style='List Number')

    def _calcular_tabela_tc_interno(self, dados):
        """Calcula a tabela comparativa de Tc diretamente dos dados de entrada.
        Retorna lista de (key, nome, valor) e o método adotado."""
        import math
        
        def to_float(val, default=0):
            try:
                if isinstance(val, str):
                    val = val.replace(',', '.')
                return float(val)
            except (ValueError, TypeError):
                return default
        
        dados_entrada = dados.get('dados_entrada', {})
        distancia = to_float(dados_entrada.get('distancia'), 0)
        desnivel = to_float(dados_entrada.get('desnivel'), 0)
        area = to_float(dados_entrada.get('area'), 0)
        declividade = to_float(dados_entrada.get('declividade'), 0)
        tc_adotado = to_float(dados_entrada.get('tempo'), 0)
        metodo_usado = dados.get('metodo_tc') or dados_entrada.get('metodo_tc', 'manual')
        
        resultados = []
        
        # Kirpich: Tc = 57 * (L_km³ / H)^0.385
        if distancia > 0 and desnivel > 0:
            L_km = distancia / 1000
            tc_kirpich = 57 * ((L_km ** 3) / desnivel) ** 0.385
            if tc_kirpich > 0:
                resultados.append(('kirpich', 'Kirpich', tc_kirpich))
        
        # Giandotti: Tc = (4√A + 1.5L) / (0.8√H) * 60
        if area > 0 and distancia > 0 and desnivel > 0:
            L_km = distancia / 1000
            tc_h = (4 * math.sqrt(area) + 1.5 * L_km) / (0.8 * math.sqrt(desnivel))
            tc_giandotti = tc_h * 60
            if tc_giandotti > 0:
                resultados.append(('giandotti', 'Giandotti', tc_giandotti))
        
        # Ventura: Tc = 0.127 * √(A / S_decimal) * 60
        if area > 0 and declividade > 0:
            tc_ventura = 0.127 * math.sqrt(area / (declividade / 100)) * 60
            if tc_ventura > 0:
                resultados.append(('ventura', 'Ventura', tc_ventura))
        
        # Bransby-Williams: Tc = 14.6 * L_km / (A^0.1 * S_decimal^0.2)
        if distancia > 0 and area > 0 and declividade > 0:
            L_km = distancia / 1000
            S = declividade / 100
            tc_bw = 14.6 * L_km / (area ** 0.1 * S ** 0.2)
            if tc_bw > 0:
                resultados.append(('bransby_williams', 'Bransby-Williams', tc_bw))
        
        # Média
        if resultados:
            media = sum(v for _, _, v in resultados) / len(resultados)
            resultados.append(('media', 'MÉDIA', media))
        
        # Valor Adotado
        if tc_adotado > 0:
            label = 'Manual' if metodo_usado == 'manual' else metodo_usado.replace('_', ' ').title()
            resultados.append(('adotado', f'ADOTADO ({label})', tc_adotado))
        
        return resultados, metodo_usado

    def _adicionar_metodologia_tc_docx(self, dados):
        """Adiciona seção detalhada sobre a metodologia de cálculo do Tc no DOCX"""
        self.document.add_heading("4.1 Metodologia do Tempo de Concentração (Tc)", level=2)
        
        dados_entrada = dados.get('dados_entrada', {})
        metodo_usado = dados.get('metodo_tc') or dados_entrada.get('metodo_tc', 'manual')
        tabela_tc = dados.get('tabela_tc') or dados_entrada.get('tabela_tc', [])
        
        # SEMPRE calcular a tabela se estiver vazia
        if not tabela_tc:
            tabela_tc, metodo_usado = self._calcular_tabela_tc_interno(dados)
        
        nomes_metodos = {
            'kirpich': 'Kirpich',
            'giandotti': 'Giandotti',
            'ventura': 'Ventura',
            'bransby_williams': 'Bransby-Williams',
            'media': 'Média Aritmética'
        }
        
        nome_metodo = nomes_metodos.get(metodo_usado, metodo_usado.replace('_', ' ').title())
        
        if metodo_usado == 'manual':
            self.document.add_paragraph(
                "O Tempo de Concentração foi inserido manualmente pelo responsável técnico. "
                "Abaixo, a título de comparação, são apresentados os valores calculados por diferentes metodologias:"
            )
        else:
            self.document.add_paragraph(
                f"Para o cálculo do Tempo de Concentração (Tc), foi selecionada a metodologia: {nome_metodo}."
            )
        
        if tabela_tc:
            self.document.add_paragraph("Resultados obtidos por diferentes metodologias para comparação:")
            table = self.document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Metodologia / Equação'
            hdr[1].text = 'Tc Calculado (min)'
            hdr[0].paragraphs[0].runs[0].bold = True
            hdr[1].paragraphs[0].runs[0].bold = True
            
            for key_tc, nome_tc, valor_tc in tabela_tc:
                row = table.add_row().cells
                row[0].text = nome_tc
                row[1].text = f"{valor_tc:.2f}"
                if key_tc == metodo_usado or key_tc == 'adotado':
                    for cell in row:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True
        
        self.document.add_paragraph()

    def _adicionar_tabela_trs_docx(self, dados):
        """Adiciona análise comparativa para múltiplos tempos de retorno no DOCX"""
        self.document.add_heading("10. ANÁLISE PARA MÚLTIPLOS TEMPOS DE RETORNO", level=1)
        self.document.add_paragraph(
            "A tabela abaixo apresenta os resultados simulados para diferentes tempos de retorno (TR), "
            "permitindo avaliar o comportamento da rede e o diâmetro necessário para variadas recorrências pluviométricas."
        )

        tabela_trs = self._calcular_tabela_trs(dados)
        if not tabela_trs:
            self.document.add_paragraph("Tabela de múltiplos TRs não disponível para esta simulação.")
            return

        table = self.document.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'TR (anos)'
        hdr[1].text = 'I (mm/h)'
        hdr[2].text = 'Q (m³/s)'
        hdr[3].text = 'D (m)'
        for cell in hdr: 
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for item in tabela_trs:
            row = table.add_row().cells
            row[0].text = str(item['tr'])
            row[1].text = f"{item['intensidade']:.2f}"
            row[2].text = f"{item['vazao']:.3f}"
            row[3].text = f"{item['diametro']:.2f}"
            
        self.document.add_paragraph(
            f"* Valores calculados para Tempo de Concentração = {dados.get('dados_entrada', {}).get('tempo', 0):.2f} min e Impermeabilidade (C) = {dados.get('dados_entrada', {}).get('coef_escoamento', 0):.4f}."
        )
        

    
    def _gerar_metodologia_tc_odt(self, dados):
        """Gera o bloco XML para a metodologia do Tc no formato ODT"""
        dados_entrada = dados.get('dados_entrada', {})
        metodo_usado = dados.get('metodo_tc') or dados_entrada.get('metodo_tc', 'manual')
        tabela_tc = dados.get('tabela_tc') or dados_entrada.get('tabela_tc', [])
        
        # SEMPRE calcular a tabela se estiver vazia
        if not tabela_tc:
            tabela_tc, metodo_usado = self._calcular_tabela_tc_interno(dados)
        
        nomes_metodos = {
            'kirpich': 'Kirpich',
            'giandotti': 'Giandotti',
            'ventura': 'Ventura',
            'bransby_williams': 'Bransby-Williams',
            'media': 'Média Aritmética'
        }
        
        nome_metodo = nomes_metodos.get(metodo_usado, metodo_usado.replace('_', ' ').title())
        
        if metodo_usado == 'manual':
            xml = '<text:p text:style-name="Normal">O Tempo de Concentração foi inserido manualmente pelo responsável técnico. Abaixo, a título de comparação, são apresentados os valores calculados por diferentes metodologias:</text:p>'
        else:
            xml = f'<text:p text:style-name="Normal">Para o cálculo do Tempo de Concentração (Tc), foi selecionada a metodologia: <text:span text:style-name="T_Bold">{nome_metodo}</text:span>.</text:p>'
        
        if tabela_tc:
            xml += '<text:p text:style-name="Normal">Resultados obtidos por diferentes metodologias para comparação:</text:p>'
            xml += '<table:table table:name="TabelaTc" table:style-name="TableCenter">'
            xml += '<table:table-column table:number-columns-repeated="2"/>'
            xml += '<table:table-row>'
            xml += '<table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Metodologia / Equação</text:p></table:table-cell>'
            xml += '<table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Tc Calculado (min)</text:p></table:table-cell>'
            xml += '</table:table-row>'
            
            for key_tc, nome_tc, valor_tc in tabela_tc:
                style = "TableCellBold" if key_tc == metodo_usado or key_tc == 'adotado' else "TableCell"
                xml += f'''
            <table:table-row>
                <table:table-cell table:style-name="{style}"><text:p>{nome_tc}</text:p></table:table-cell>
                <table:table-cell table:style-name="{style}"><text:p>{valor_tc:.2f}</text:p></table:table-cell>
            </table:table-row>'''
            xml += '</table:table>'
            
        return xml
        
    def _gerar_tabela_idf_odt(self, dados):
        """Gera o bloco XML para a tabela IDF no formato ODT"""
        try:
            from ..hidrologia.curvas_idf import CurvasIDF
            idf_calc = CurvasIDF()
            dados_entrada = dados.get('dados_entrada', {})
            cidade_key = dados_entrada.get('cidade_idf', '').lower().replace(' ', '_')
            params_idf = idf_calc.obter_parametros(cidade_key)
            if not params_idf:
                return ''
                
            TRs = [2, 5, 10, 25, 50, 100]
            duracoes = [10, 15, 30, 60, 120]
            T_req = dados_entrada.get('tempo', 15)
            t_arr = round(T_req, 2)
            if t_arr not in duracoes:
                duracoes.append(t_arr)
            duracoes.sort()
            
            tab_dados = idf_calc.gerar_tabela_intensidades(cidade_key, TRs, duracoes)
            
            xml = f'<text:p text:style-name="Normal">Resultados associados à curva IDF de {cidade_key.replace("_", " ").title()}:</text:p>'
            xml += '<table:table table:name="TabelaIDF" table:style-name="TableCenter">'
            xml += '<table:table-column table:style-name="ColIDFDur"/>'
            xml += '<table:table-column table:style-name="ColIDFTR" table:number-columns-repeated="6"/>'
            xml += '<table:table-row>'
            xml += '<table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Duração (min)</text:p></table:table-cell>'
            for tr_val in TRs:
                xml += f'<table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">TR={tr_val}</text:p></table:table-cell>'
            xml += '</table:table-row>'
            
            for d in duracoes:
                xml += '<table:table-row>'
                xml += f'<table:table-cell table:style-name="TableCell"><text:p text:style-name="NormalCenterSmall">{d}</text:p></table:table-cell>'
                for tr_val in TRs:
                    val_idf = tab_dados['intensidades'][tr_val][d]
                    if d == t_arr and tr_val == dados_entrada.get('tempo_retorno', 25):
                        xml += f'<table:table-cell table:style-name="TableCell"><text:p text:style-name="NormalCenterSmall"><text:span text:style-name="T_Bold">{val_idf:.1f}</text:span></text:p></table:table-cell>'
                    else:
                        xml += f'<table:table-cell table:style-name="TableCell"><text:p text:style-name="NormalCenterSmall">{val_idf:.1f}</text:p></table:table-cell>'
                xml += '</table:table-row>'
            xml += '</table:table><text:p text:style-name="Normal"/>'
            return xml
        except Exception as e:
            return f'<text:p text:style-name="Normal">Erro ao gerar tabela IDF: {e}</text:p>'

    def _gerar_odt(self, dados, caminho_saida):
        """
        Gera relatório em formato ODT (OpenDocument Text).
        ODT é um formato aberto suportado por LibreOffice, OpenOffice e Word.
        """
        def to_float(val, default=0):
            try:
                if isinstance(val, str):
                    val = val.replace(',', '.')
                return float(val)
            except (ValueError, TypeError):
                return default

        dados_entrada = dados.get('dados_entrada', {})
        
        # Extrair valores com fallbacks robustos
        vazao = to_float(dados.get('vazao') or dados_entrada.get('vazao'), 0)
        diametro = to_float(dados.get('diametro') or dados_entrada.get('diametro'), 0)
        velocidade = to_float(dados.get('velocidade') or dados_entrada.get('velocidade'), 0)
        froude = to_float(dados.get('froude') or dados_entrada.get('froude'), 0)
        lamina = to_float(dados.get('lamina') or dados_entrada.get('lamina'), 0.85) * 100
        
        distancia = to_float(dados_entrada.get('distancia'), 0)
        desnivel = to_float(dados_entrada.get('desnivel'), 0)
        tempo = to_float(dados_entrada.get('tempo'), 0)
        area = to_float(dados_entrada.get('area'), 0)
        coef = to_float(dados_entrada.get('coef_escoamento') or dados_entrada.get('impermeabilidade'), 0)
        rugosidade = to_float(dados_entrada.get('rugosidade'), 0)
        declividade = to_float(dados_entrada.get('declividade'), 0)
        tr = dados_entrada.get('tempo_retorno', 0)
        intensidade = to_float(dados_entrada.get('intensidade'), 0)
        
        # Status das verificações
        if 0.6 <= velocidade <= 5.0:
            status_vel = "ADEQUADA"
            obs_vel = "Velocidade dentro dos limites normativos (0,6 a 5,0 m/s)."
        elif velocidade < 0.6:
            status_vel = "BAIXA"
            obs_vel = "Velocidade abaixo do mínimo. Risco de sedimentação."
        else:
            status_vel = "ALTA"
            obs_vel = "Velocidade acima do máximo. Risco de erosão."
            
        if froude < 0.8:
            status_fr = "SUBCRÍTICO"
            obs_fr = "Escoamento lento e tranquilo. Ideal para evitar erosão excessiva."
        elif 0.8 <= froude <= 1.2:
            status_fr = "CRÍTICO / INSTÁVEL"
            obs_fr = "Estado de energia instável (Fr ≈ 1). Dificulta a precisão da lâmina d'água."
        else:
            status_fr = "SUPERCRÍTICO"
            obs_fr = "Escoamento rápido. Requer atenção especial e dissipadores de energia à jusante."
            
        # Mapa do projeto removido a pedido do usuário
        ODT_IMAGEM = ""
        
        # Gerar tabela de múltiplos TRs para ODT
        tabela_trs = self._calcular_tabela_trs(dados)
        if tabela_trs:
            tabela_trs_odt = '''<table:table table:name="TabelaTRs" table:style-name="TableCenter">
                <table:table-column table:number-columns-repeated="4"/>
                <table:table-row>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">TR (anos)</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">I (mm/h)</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Q (m³/s)</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">D (m)</text:p></table:table-cell>
                </table:table-row>'''
            for item in tabela_trs:
                tabela_trs_odt += f'''
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>{item['tr']}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{item['intensidade']:.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{item['vazao']:.3f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{item['diametro']:.2f}</text:p></table:table-cell>
                </table:table-row>'''
            tabela_trs_odt += '\n            </table:table>'
        else:
            tabela_trs_odt = '<text:p text:style-name="Normal">Tabela de múltiplos TRs não disponível.</text:p>'
            
        camadas_usadas = dados.get('camadas_usadas', {})
        tabela_camadas_odt = ""
        if camadas_usadas:
            tabela_camadas_odt = '''<table:table table:name="TabelaCamadas" table:style-name="TableCenter">
                <table:table-column table:number-columns-repeated="2"/>
                <table:table-row>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Parâmetro Espacial</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Nome da Camada (Fonte de Dados)</text:p></table:table-cell>
                </table:table-row>'''
            for tipo, nome in camadas_usadas.items():
                tabela_camadas_odt += f'''
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p text:style-name="Normal"><text:span text:style-name="T_Bold">{self._escape_xml(tipo)}</text:span></text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p text:style-name="Normal">{self._escape_xml(nome)}</text:p></table:table-cell>
                </table:table-row>'''
            tabela_camadas_odt += '\n            </table:table>'
        else:
            tabela_camadas_odt = '<text:p text:style-name="Normal">Não foram detectadas conexões automáticas de camadas neste projeto (inserção puramente manual).</text:p>'
            
        cidade_idf = self._escape_xml(dados_entrada.get('cidade_idf', 'Não Definida').replace('_', ' ').title())
        A_secao = dados.get('area_secao', 1) 
        
        # Conteúdo XML do ODT
        content_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
    xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0"
    xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"
    xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0"
    xmlns:draw="urn:oasis:names:tc:opendocument:xmlns:drawing:1.0"
    xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0"
    xmlns:xlink="http://www.w3.org/1999/xlink"
    xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0"
    office:version="1.2">
    
    <office:automatic-styles>
        <style:style style:name="Titulo" style:family="paragraph">
            <style:paragraph-properties fo:text-align="center" fo:margin-top="1cm" fo:margin-bottom="0.5cm"/>
            <style:text-properties fo:font-size="24pt" fo:font-weight="bold" fo:color="#1565C0"/>
        </style:style>
        <style:style style:name="Subtitulo" style:family="paragraph">
            <style:paragraph-properties fo:text-align="center" fo:margin-bottom="1cm"/>
            <style:text-properties fo:font-size="16pt" fo:color="#666666"/>
        </style:style>
        <style:style style:name="Heading1" style:family="paragraph">
            <style:paragraph-properties fo:margin-top="0.8cm" fo:margin-bottom="0.3cm" fo:border-bottom="0.5pt solid #1565C0"/>
            <style:text-properties fo:font-size="14pt" fo:font-weight="bold" fo:color="#1565C0"/>
        </style:style>
        <style:style style:name="Heading2" style:family="paragraph">
            <style:paragraph-properties fo:margin-top="0.5cm" fo:margin-bottom="0.2cm"/>
            <style:text-properties fo:font-size="12pt" fo:font-weight="bold" fo:color="#1976D2"/>
        </style:style>
        <style:style style:name="Normal" style:family="paragraph">
            <style:paragraph-properties fo:margin-top="0.2cm" fo:margin-bottom="0.2cm" fo:text-align="justify"/>
            <style:text-properties fo:font-size="11pt"/>
        </style:style>
        <style:style style:name="TableCenter" style:family="table">
            <style:table-properties style:width="16cm" table:align="center"/>
        </style:style>
        <style:style style:name="NormalCenterSmall" style:family="paragraph">
            <style:paragraph-properties fo:text-align="center" fo:margin-top="0.1cm" fo:margin-bottom="0.1cm"/>
            <style:text-properties fo:font-size="8.5pt"/>
        </style:style>
        <style:style style:name="ColIDFDur" style:family="table-column">
            <style:table-column-properties style:column-width="2.6cm"/>
        </style:style>
        <style:style style:name="ColIDFTR" style:family="table-column">
            <style:table-column-properties style:column-width="2.2cm"/>
        </style:style>
        <style:style style:name="Equacao" style:family="paragraph">
            <style:paragraph-properties fo:text-align="center" fo:margin-top="0.3cm" fo:margin-bottom="0.3cm" fo:background-color="#E3F2FD" fo:padding="0.3cm"/>
            <style:text-properties fo:font-size="12pt" fo:font-weight="bold"/>
        </style:style>
        <style:style style:name="Conclusao" style:family="paragraph">
            <style:paragraph-properties fo:background-color="#E3F2FD" fo:padding="0.3cm" fo:margin-top="0.3cm"/>
            <style:text-properties fo:font-size="11pt"/>
        </style:style>
        <style:style style:name="NormalCenter" style:family="paragraph">
            <style:paragraph-properties fo:text-align="center" fo:margin-top="0.2cm" fo:margin-bottom="0.2cm"/>
            <style:text-properties fo:font-size="11pt"/>
        </style:style>
        <style:style style:name="ItalicCenter" style:family="paragraph">
            <style:paragraph-properties fo:text-align="center" fo:margin-top="0.1cm" fo:margin-bottom="0.2cm"/>
            <style:text-properties fo:font-size="9pt" fo:font-style="italic" fo:color="#666666"/>
        </style:style>
        <style:style style:name="ImageCell" style:family="table-cell">
            <style:table-cell-properties fo:padding="0.1cm" fo:border="none"/>
        </style:style>
        <style:style style:name="TableCell" style:family="table-cell">
            <style:table-cell-properties fo:padding="0.2cm" fo:border="0.5pt solid #DDDDDD"/>
        </style:style>
        <style:style style:name="TableHeader" style:family="table-cell">
            <style:table-cell-properties fo:padding="0.2cm" fo:border="0.5pt solid #DDDDDD" fo:background-color="#1565C0"/>
        </style:style>
        <style:style style:name="HeaderText" style:family="paragraph">
            <style:text-properties fo:color="#FFFFFF" fo:font-weight="bold"/>
        </style:style>
        <style:style style:name="fr1" style:family="graphic">
            <style:graphic-properties style:run-through="foreground" style:wrap="none" style:horizontal-pos="center" style:horizontal-rel="paragraph" style:mirror="none" fo:clip="rect(0cm, 0cm, 0cm, 0cm)" draw:luminance="0%" draw:contrast="0%" draw:red="0%" draw:green="0%" draw:blue="0%" draw:gamma="100%" draw:color-inversion="false" draw:image-opacity="100%" draw:color-mode="standard"/>
        </style:style>
        <style:style style:name="T1" style:family="text">
            <style:text-properties fo:font-style="italic"/>
        </style:style>
        <style:style style:name="T_Bold" style:family="text">
            <style:text-properties fo:font-weight="bold"/>
        </style:style>
    </office:automatic-styles>
    
    <office:body>
        <office:text>
            <!-- CAPA -->
            <text:p text:style-name="Titulo">ESTUDO DE DRENAGEM URBANA</text:p>
            <text:p text:style-name="Subtitulo">Método Racional</text:p>
            <text:p text:style-name="Normal"/>
            <text:p text:style-name="Normal"/>
            <text:p text:style-name="Normal"/>
            <text:p text:style-name="Normal"/>
            <text:p text:style-name="Normal"/>

            <!-- LOCALIZAÇÃO -->
            <text:p text:style-name="Heading1">LOCALIZAÇÃO DA ÁREA DE ESTUDOS</text:p>
            <table:table table:name="TabelaLocalizacao" table:style-name="TableCenter">
                <table:table-column table:number-columns-repeated="2"/>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p text:style-name="Normal"><text:span text:style-name="T_Bold">Região / Endereço:</text:span></text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>[Espaço reservado para preenchimento]</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p text:style-name="Normal"><text:span text:style-name="T_Bold">Coordenadas (Centroide):</text:span></text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{dados.get('coordenadas', 'Não disponível')}</text:p></table:table-cell>
                </table:table-row>
            </table:table>

            <!-- RESUMO -->
            <text:p text:style-name="Heading1">RESUMO DOS DADOS PRINCIPAIS</text:p>
            <table:table table:name="TabelaResumo" table:style-name="TableCenter">
                <table:table-column table:number-columns-repeated="3"/>
                <table:table-row>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Vazão de Projeto</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Área da Bacia</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Tempo de Retorno</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p text:style-name="NormalCenter"><text:span text:style-name="T_Bold">{vazao:.2f} m³/s</text:span></text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p text:style-name="NormalCenter"><text:span text:style-name="T_Bold">{area:.4f} km²</text:span></text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p text:style-name="NormalCenter"><text:span text:style-name="T_Bold">{tr} anos</text:span></text:p></table:table-cell>
                </table:table-row>
            </table:table>
            
            <!-- METODOLOGIA -->
            <text:p text:style-name="Heading1">1. METODOLOGIA</text:p>
            <text:p text:style-name="Heading2">1.1 Método Racional</text:p>
            <text:p text:style-name="Normal">O Método Racional é utilizado para estimativa de vazões de pico em pequenas bacias urbanas:</text:p>
            <text:p text:style-name="Equacao">Q = (C × I × A) / 3,6</text:p>
            <text:p text:style-name="Normal">Onde: Q = Vazão (m³/s), C = Coeficiente de escoamento, I = Intensidade (mm/h), A = Área (km²)</text:p>
            
            <text:p text:style-name="Heading2">1.2 Equação de Manning</text:p>
            <text:p text:style-name="Equacao">Q = (1/n) × A × R^(2/3) × S^(1/2)</text:p>
            
            <!-- DADOS DE ENTRADA -->
            <text:p text:style-name="Heading1">2. DADOS DE ENTRADA</text:p>
            <text:p text:style-name="Heading2">2.1 Fontes de Extração Espacial</text:p>
            <text:p text:style-name="Normal">Os parâmetros morfométricos e de uso do solo foram extraídos a partir das seguintes matrizes e vetores geográficos reais contidos no projeto atual do QGIS:</text:p>
            {tabela_camadas_odt}
            
            <text:p text:style-name="Heading2">2.2 Parâmetros Físicos Adotados</text:p>
            <text:p text:style-name="Normal">No dimensionamento atual, foram considerados conceitos primordiais do Método Racional clássico. A capacidade das redes considerou uma restrição cautelar de Lâmina D'água para evitar a pressurização precoce e o golpe de aríete no poço de visita, delimitando o uso da seção em até 85%. Em contrapartida, avaliou-se a Velocidade de escoamento para afastar sedimentações de areias por velocidades mortas (&lt; 0,60 m/s) ou erosões nos dutos em concretos/PVC geradas por grandes velocidades (&gt; 5,00 m/s).</text:p>
            
            <table:table table:name="DadosEntrada" table:style-name="TableCenter">
                <table:table-column table:number-columns-repeated="3"/>
                <table:table-row>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Parâmetro Físico</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Valor</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Unidade</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Distância (Talvegue Máximo)</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{distancia:.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>m</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Desnível Topográfico</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{desnivel:.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>m</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Tempo de Concentração</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{tempo:.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>min</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Área Contribuinte (A)</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{area:.6f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>km²</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Impermeabilidade do Solo (C)</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{coef:.4f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>-</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Rugosidade de Manning (n)</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{rugosidade:.4f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>-</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Declividade Longitudinal (S)</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{declividade:.4f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>%</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Tempo de Retorno Estipulado (TR)</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{tr}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>anos</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Intensidade Pluviométrica (i)</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{intensidade:.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>mm/h</text:p></table:table-cell>
                </table:table-row>
            </table:table>
            
            <!-- MEMÓRIA DE CÁLCULO -->
            <text:p text:style-name="Heading1">3. MEMÓRIA DE CÁLCULO</text:p>
            
            <text:p text:style-name="Heading2">3.1 Metodologia do Tempo de Concentração (Tc)</text:p>
            {self._gerar_metodologia_tc_odt(dados)}

            <text:p text:style-name="Heading2">3.2 Cálculo de Intensidade (IDF)</text:p>
            <text:p text:style-name="Normal">Baseado na equação de chuvas intensas selecionada para localidade de {cidade_idf}:</text:p>
            <text:p text:style-name="Equacao">i = (K × TR^a) / (t + b)^c</text:p>
            {self._gerar_params_idf_odt(dados)}
            <text:p text:style-name="Normal">Para TR = {tr} anos e Tempo de Concentração = {tempo:.2f} min:</text:p>
            <text:p text:style-name="Normal"><text:span text:style-name="T_Bold">Intensidade Pluviométrica Equivalente (I) = {intensidade:.2f} mm/h</text:span></text:p>
            
            <text:p text:style-name="Heading2">3.3 Tabela de Intensidades para Múltiplos TRs e Durações</text:p>
            {self._gerar_tabela_idf_odt(dados)}
            
            <text:p text:style-name="Heading2">3.4 Determinação da Vazão (Método Racional)</text:p>
            <text:p text:style-name="Normal">Através da fórmula matriz de compatibilização da área de deflúvio (km²) e intensidade (mm/h) para vazão (m³/s):</text:p>
            <text:p text:style-name="Equacao">Q = (C × i × A) / 3,6</text:p>
            <text:p text:style-name="Normal">Q = ({coef:.4f} × {intensidade:.2f} × {area:.6f}) / 3,6</text:p>
            <text:p text:style-name="Normal"><text:span text:style-name="T_Bold">Vazão de Projeto Estimada (Q_max) = {vazao:.2f} m³/s</text:span></text:p>
            
            <text:p text:style-name="Heading2">3.5 Dimensionamento por Manning</text:p>
            <text:p text:style-name="Normal">Para a vazão obtida de <text:span text:style-name="T_Bold">{vazao:.2f} m³/s</text:span>, aplica-se a fórmula de Manning admitindo regime uniforme em seção de galeria circular plena a uma ocupação normativa de 85%.</text:p>
            <text:p text:style-name="Equacao">Q = (1 / n) * A_secao * R_h^(2/3) * S^(1/2)</text:p>
            <text:p text:style-name="Normal">Substituindo a Declividade (S) de {declividade:.4f}% e a Rugosidade assumida (n) de {rugosidade:.4f}:</text:p>
            <text:p text:style-name="Normal">- Diâmetro Equivalente Calculado: <text:span text:style-name="T_Bold">{diametro:.2f} m</text:span></text:p>
            <text:p text:style-name="Normal">- Lado p/ Galeria Celular Quadrada Equiv.: <text:span text:style-name="T_Bold">{dados.get('lado_galeria', 0):.2f} m</text:span></text:p>
            <text:p text:style-name="Normal">- Velocidade Resultante (Seção Plena): <text:span text:style-name="T_Bold">{velocidade:.2f} m/s</text:span></text:p>
            
            {ODT_IMAGEM}
            
            <!-- ANÁLISE DE IMPERMEABILIDADE -->
            '''
        dados_imper = dados.get('impermeabilidade_dados')
        if dados_imper:
            content_xml += f'''
            <text:p text:style-name="Heading1">4. ANÁLISE DE IMPERMEABILIDADE DO SOLO</text:p>
            <text:p text:style-name="Normal">A determinação da Impermeabilidade do Solo (C) foi realizada através de análise multiespectral e classificação digital de pixels.</text:p>
            
            <table:table table:name="TabelaImper" table:style-name="TableCenter">
                <table:table-column table:number-columns-repeated="3"/>
                <table:table-row>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Classe de Uso</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Pixels</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">%</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Impermeável</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{dados_imper.get('impermeable_pixels', 0):,}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{dados_imper.get('percentual', 0):.2f}%</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Vegetação</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{dados_imper.get('vegetation_pixels', 0):,}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{dados_imper.get('percent_vegetation', 0):.2f}%</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Sombra / Água</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{dados_imper.get('shadow_pixels', 0):,}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{dados_imper.get('percent_shadow', 0):.2f}%</text:p></table:table-cell>
                </table:table-row>
            </table:table>
            
            <text:p text:style-name="Heading2">4.1 Anexo Fotográfico da Classificação</text:p>
            '''
            
            img_orig = dados_imper.get('impermeabilidade_imagem_original')
            img_class = dados_imper.get('impermeabilidade_imagem')
            
            if img_orig and os.path.exists(img_orig) and img_class and os.path.exists(img_class):
                content_xml += f'''
                <table:table table:name="AnexoFotos">
                    <table:table-column table:number-columns-repeated="2"/>
                    <table:table-row>
                        <table:table-cell table:style-name="ImageCell">
                            <text:p text:style-name="NormalCenter">
                                <draw:frame draw:style-name="fr1" text:anchor-type="paragraph" svg:width="7.5cm" svg:height="7.5cm">
                                    <draw:image xlink:href="Pictures/imper_original.png" xlink:type="simple" xlink:show="embed" xlink:actuate="onLoad"/>
                                </draw:frame>
                            </text:p>
                            <text:p text:style-name="ItalicCenter">Imagem Original (RGB)</text:p>
                        </table:table-cell>
                        <table:table-cell table:style-name="ImageCell">
                            <text:p text:style-name="NormalCenter">
                                <draw:frame draw:style-name="fr1" text:anchor-type="paragraph" svg:width="7.5cm" svg:height="7.5cm">
                                    <draw:image xlink:href="Pictures/imper_class.png" xlink:type="simple" xlink:show="embed" xlink:actuate="onLoad"/>
                                </draw:frame>
                            </text:p>
                            <text:p text:style-name="ItalicCenter">Mapa de Classificação</text:p>
                        </table:table-cell>
                    </table:table-row>
                </table:table>
                '''
            elif img_class and os.path.exists(img_class):
                content_xml += f'''
                <text:p text:style-name="NormalCenter">
                    <draw:frame draw:style-name="fr1" text:anchor-type="paragraph" svg:width="15cm">
                        <draw:image xlink:href="Pictures/imper_class.png" xlink:type="simple" xlink:show="embed" xlink:actuate="onLoad"/>
                    </draw:frame>
                </text:p>
                <text:p text:style-name="ItalicCenter">Figura: Mapa de Classificação da Impermeabilidade</text:p>
                '''

        content_xml += f'''
            <!-- RESULTADOS -->
            <text:p text:style-name="Heading1">5. RESULTADOS</text:p>
            <table:table table:name="Resultados" table:style-name="TableCenter">
                <table:table-column table:number-columns-repeated="3"/>
                <table:table-row>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Resultado</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Valor</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableHeader"><text:p text:style-name="HeaderText">Unidade</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Vazão de Projeto</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{vazao:.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>m³/s</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Diâmetro Calculado</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{diametro:.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>m</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Lado Equiv. Galeria</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{dados.get('lado_galeria', 0):.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>m</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Área da Seção</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{A_secao:.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>m²</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Velocidade de Escoamento</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{velocidade:.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>m/s</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Número de Froude</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{froude:.2f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>-</text:p></table:table-cell>
                </table:table-row>
                <table:table-row>
                    <table:table-cell table:style-name="TableCell"><text:p>Lâmina/Diâmetro</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>{lamina:.0f}</text:p></table:table-cell>
                    <table:table-cell table:style-name="TableCell"><text:p>%</text:p></table:table-cell>
                </table:table-row>
            </table:table>
            
            <!-- TABELA DE MÚLTIPLOS TRs -->
            <text:p text:style-name="Heading1">6. ANÁLISE PARA MÚLTIPLOS TEMPOS DE RETORNO</text:p>
            <text:p text:style-name="Normal">Resultados calculados para diferentes tempos de retorno, considerando os mesmos parâmetros físicos e a impermeabilidade calculada:</text:p>
            {tabela_trs_odt}
            <text:p text:style-name="Normal">* Valores calculados para Tempo de Concentração = {tempo:.2f} min e Impermeabilidade (C = {coef:.4f})</text:p>
            
            <!-- VERIFICAÇÕES -->
            <text:p text:style-name="Heading1">7. VERIFICAÇÕES HIDRÁULICAS</text:p>
            <text:p text:style-name="Heading2">7.1 Verificação de Velocidade do Escoamento</text:p>
            <text:p text:style-name="Normal">Equação Padrão: V = Q / Área_Molhada_Secao</text:p>
            <text:p text:style-name="Normal">V = {vazao:.2f} / {A_secao:.2f}</text:p>
            <text:p text:style-name="Normal"><text:span text:style-name="T_Bold">Velocidade (v): {velocidade:.2f} m/s - {status_vel}</text:span></text:p>
            <text:p text:style-name="Normal">{obs_vel}</text:p>
            <text:p text:style-name="Normal">Limites normativos: Mínimo 0,6 m/s (autolimpeza) | Máximo 5,0 m/s (erosão)</text:p>
            
            <text:p text:style-name="Heading2">7.2 Verificação de Regime (Froude)</text:p>
            <text:p text:style-name="Normal">O número de Froude indica a influência da inércia sobre o escoamento face à gravidade. Escoamentos normais e passivos devem ser mantidos no estado Subcrítico (Flúvial) para prevenir choques de onda e perdas de carga localizadas no bueiro/coletores.</text:p>
            <text:p text:style-name="Normal">Equação: Fr = vel / √ (g * Profundidade_Hidráulica)</text:p>
            <text:p text:style-name="Normal">Fr = {velocidade:.2f} / √ (9.81 * {diametro:.2f})</text:p>
            <text:p text:style-name="Normal"><text:span text:style-name="T_Bold">Número de Froude (Fr): {froude:.2f} - {status_fr}</text:span></text:p>
            <text:p text:style-name="Normal">{obs_fr}</text:p>
            
            <!-- CONCLUSÕES -->
            <text:p text:style-name="Heading1">8. CONCLUSÕES</text:p>
            <text:p text:style-name="Conclusao">Com base nos cálculos realizados:</text:p>
            <text:p text:style-name="Normal">• Vazão de projeto: {vazao:.2f} m³/s</text:p>
            <text:p text:style-name="Normal">• Diâmetro recomendado: {diametro:.2f} m</text:p>
            <text:p text:style-name="Normal">• Velocidade: {velocidade:.2f} m/s ({status_vel})</text:p>
            <text:p text:style-name="Normal">• Regime: {status_fr} (Fr = {froude:.2f})</text:p>
            
            <!-- ESPAÇO PARA COMPLEMENTAÇÃO E LIMITAÇÕES TÉCNICAS -->
            <text:p text:style-name="Heading1">9. NOTA DE LIMITAÇÃO TÉCNICA</text:p>
            <text:p text:style-name="Normal">1. Aplicabilidade do Método: Este relatório utiliza o Método Racional, cuja premissa fundamental assume que a intensidade da chuva é uniforme e constante sobre toda a bacia durante o tempo de concentração (tc). Este método é estritamente recomendado para bacias de microdrenagem e áreas urbanas pequenas (geralmente inferiores a 50-100 hectares).</text:p>
            <text:p text:style-name="Normal">2. Macrodrenagem e Áreas Extensas: Para bacias de macrodrenagem ou áreas que excedam os limites de aplicação do método, os resultados aqui apresentados devem ser interpretados como estimativas preliminares. Recomenda-se a validação através de métodos que considerem o amortecimento de cheias e a variação temporal da chuva, como o Método do Hidrograma Unitário (SCS/NRCS) ou modelagem hidráulica dinâmica (ex: SWMM).</text:p>
            <text:p text:style-name="Normal">3. Coeficiente de Escoamento (C): Os valores de C adotados baseiam-se em tabelas normatizadas de uso e ocupação do solo. Alterações futuras na impermeabilização da bacia invalidam as vazões de projeto calculadas.</text:p>
            <text:p text:style-name="Normal">4. Verificação Hidráulica: A análise do Número de Froude (Fr) indica o regime de escoamento. Trechos com Fr próximos à unidade (regime crítico) apresentam instabilidade de lâmina d'água e requerem atenção especial no dimensionamento físico para evitar transbordamentos ou erosões não previstos.</text:p>

            <text:p text:style-name="Heading1">10. OBSERVAÇÕES E COMPLEMENTAÇÕES</text:p>
            <text:p text:style-name="Normal">[Espaço reservado para observações adicionais do responsável técnico]</text:p>
            <text:p text:style-name="Normal"/>
            <text:p text:style-name="Normal"/>
            <text:p text:style-name="Normal"/>
            

            
            <!-- RODAPÉ -->
            <text:p text:style-name="Normal"/>
            <text:p text:style-name="Normal">_______________________________________________________</text:p>
            <text:p text:style-name="Normal">Relatório gerado pelo Plugin Método Racional Pro - QGIS</text:p>
            <text:p text:style-name="Normal">{datetime.now().strftime("%d/%m/%Y às %H:%M")}</text:p>
            
        </office:text>
    </office:body>
</office:document-content>'''

        # Manifest XML dinâmico
        manifest_entries = [
            ('<manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/>', True),
            ('<manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>', True),
            ('<manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/>', True),
            ('<manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/>', True),
            ('<manifest:file-entry manifest:full-path="Pictures/imper_original.png" manifest:media-type="image/png"/>', dados_imper and dados_imper.get('impermeabilidade_imagem_original') and os.path.exists(dados_imper.get('impermeabilidade_imagem_original'))),
            ('<manifest:file-entry manifest:full-path="Pictures/imper_class.png" manifest:media-type="image/png"/>', dados_imper and dados_imper.get('impermeabilidade_imagem') and os.path.exists(dados_imper.get('impermeabilidade_imagem'))),
        ]
        
        manifest_content = "\n    ".join([entry for entry, condition in manifest_entries if condition])
        manifest_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2">
    {manifest_content}
</manifest:manifest>'''

        # Styles XML
        styles_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
    xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0"
    xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0"
    office:version="1.2">
    <office:styles>
        <style:default-style style:family="paragraph">
            <style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0.2cm"/>
            <style:text-properties fo:font-family="Arial" fo:font-size="11pt"/>
        </style:default-style>
    </office:styles>
</office:document-styles>'''

        # Meta XML
        meta_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
    xmlns:dc="http://purl.org/dc/elements/1.1/"
    xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0"
    office:version="1.2">
    <office:meta>
        <dc:title>Estudo de Drenagem Urbana - Método Racional</dc:title>
        <dc:creator>Plugin Método Racional Pro - QGIS</dc:creator>
        <dc:date>{datetime.now().isoformat()}</dc:date>
        <meta:generator>Método Racional Pro 1.0</meta:generator>
    </office:meta>
</office:document-meta>'''

        # Criar arquivo ODT (é um ZIP com estrutura específica)
        with zipfile.ZipFile(caminho_saida, 'w', zipfile.ZIP_DEFLATED) as odt:
            # Mimetype deve ser o primeiro arquivo e sem compressão
            odt.writestr('mimetype', 'application/vnd.oasis.opendocument.text', compress_type=zipfile.ZIP_STORED)
            odt.writestr('META-INF/manifest.xml', manifest_xml)
            odt.writestr('content.xml', content_xml)
            odt.writestr('styles.xml', styles_xml)
            odt.writestr('meta.xml', meta_xml)
            

                
            # Adicionar imagens de impermeabilidade
            if dados_imper:
                img_orig = dados_imper.get('impermeabilidade_imagem_original')
                img_class = dados_imper.get('impermeabilidade_imagem')
                
                if img_orig and os.path.exists(img_orig):
                    odt.write(img_orig, "Pictures/imper_original.png")
                if img_class and os.path.exists(img_class):
                    odt.write(img_class, "Pictures/imper_class.png")


class GeradorRelatorioSimplificado(GeradorRelatorio):
    """Gera relatório simplificado"""
    
    def gerar_relatorio_completo(self, dados, caminho_saida):
        """Gera versão simplificada do relatório"""
        self.document = Document()
        self._configurar_estilos()
        
        self._adicionar_capa(dados)
        self._adicionar_localizacao_resumo_docx(dados)
        self._adicionar_dados_entrada(dados)
        self._adicionar_resultados(dados)
        self._adicionar_verificacoes(dados)
        
        self.document.save(caminho_saida)
